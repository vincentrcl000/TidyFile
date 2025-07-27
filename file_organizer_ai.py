#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI智能分类专用文件整理核心模块
"""
import os
import shutil
import logging
import json
import re
from datetime import datetime
from typing import List, Dict, Tuple, Optional, Any
from pathlib import Path
import ollama
from transfer_log_manager import TransferLogManager
from classification_rules_manager import ClassificationRulesManager
import PyPDF2
from PyPDF2 import PdfReader, PdfWriter
import docx
from docx import Document
import time
import requests
from ai_client_manager import get_ai_manager, chat_with_ai, refresh_ai_clients

class FileOrganizerError(Exception):
    pass

class FileOrganizer:
    def __init__(self, model_name: str = None, enable_transfer_log: bool = True):
        self.model_name = model_name
        self.enable_transfer_log = enable_transfer_log
        self.transfer_log_manager = None
        self.setup_logging()
        
        if self.enable_transfer_log:
            self.transfer_log_manager = TransferLogManager()
        
        # 初始化分类规则管理器
        self.classification_rules_manager = ClassificationRulesManager()
        
        # 初始化AI管理器
        self.ai_manager = get_ai_manager()
        
        # 初始化AI参数
        self.ai_parameters = {
            'similarity_threshold': 0.7,
            'max_retries': 3,
            'content_extraction_length': 3000,
            'summary_length': 200,
            'classification_prompt_template': None
        }
        
        # 新增：文件缓存和标签系统
        self.file_cache = {}  # 缓存文件信息和元数据
        self.level_tags = {}  # 缓存各级标签
        self.global_cache = {}  # 全局缓存，用于总结和删除源文件等功能
    def setup_logging(self) -> None:
        """设置日志配置，仅输出到控制台"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.StreamHandler()
            ]
        )
        logging.info("文件整理器初始化完成")
    def initialize_ollama(self) -> None:
        """初始化AI客户端，使用统一的AI管理器"""
        try:
            # 刷新AI客户端
            refresh_ai_clients()
            logging.info("AI客户端初始化成功")
        except Exception as e:
            raise FileOrganizerError(f"AI客户端初始化失败: {e}")
    def scan_target_folders(self, target_directory: str) -> List[str]:
        """扫描目标文件夹，返回相对路径列表（内部使用，不输出日志）"""
        try:
            target_path = Path(target_directory)
            if not target_path.exists():
                raise FileOrganizerError(f"目标目录不存在: {target_directory}")
            folders = []
            for item in target_path.rglob('*'):
                if item.is_dir() and item != target_path:  # 排除目标目录本身
                    relative_path = item.relative_to(target_path)
                    folders.append(str(relative_path))
            return folders
        except Exception as e:
            raise FileOrganizerError(f"扫描目标文件夹失败: {e}")
    def get_directory_tree_structure(self, target_directory: str) -> str:
        """生成目标目录结构，返回完整路径列表"""
        try:
            target_path = Path(target_directory)
            if not target_path.exists():
                raise FileOrganizerError(f"目标目录不存在: {target_directory}")
            
            # 直接扫描并构建目录结构
            folders = []
            for item in target_path.rglob('*'):
                if item.is_dir() and item != target_path:  # 排除目标目录本身
                    relative_path = item.relative_to(target_path)
                    folders.append(str(relative_path))
            
            # 构建清晰的路径列表，每行一个完整路径
            path_lines = []
            for i, folder_path in enumerate(folders, 1):
                path_lines.append(f"{i}. {folder_path}")
            
            tree_structure = "\n".join(path_lines)
            logging.info(f"生成目录路径结构，共 {len(path_lines)} 行")
            return tree_structure
        except Exception as e:
            raise FileOrganizerError(f"生成目录树结构失败: {e}")
    def analyze_and_classify_file(self, file_path: str, target_directory: str) -> Dict[str, Any]:
        """
        新的文件分析和分类方法：先解析内容，生成摘要，再推荐目录
        返回包含提取内容、摘要和推荐目录的完整结果
        """
        start_time = time.time()
        timing_info = {}
        
        try:
            if not hasattr(self, 'ai_manager') or self.ai_manager is None:
                init_start = time.time()
                self.initialize_ollama()
                timing_info['ollama_init_time'] = round(time.time() - init_start, 3)
            
            file_name = Path(file_path).name
            logging.info(f"开始分析文件: {file_name}")
            
            # 第一步：提取文件元数据并缓存
            metadata_start = time.time()
            file_metadata = self._extract_file_metadata(file_path)
            self.file_cache[file_path] = file_metadata
            metadata_time = round(time.time() - metadata_start, 3)
            timing_info['metadata_extraction_time'] = metadata_time
            logging.info(f"文件元数据提取完成，耗时: {metadata_time}秒")
            
            # 第二步：解析文件内容
            extract_start = time.time()
            extracted_content = self._extract_file_content(file_path)
            extract_time = round(time.time() - extract_start, 3)
            timing_info['content_extraction_time'] = extract_time
            logging.info(f"文件内容提取完成，长度: {len(extracted_content)} 字符，耗时: {extract_time}秒")
            
            # 根据设置截取内容用于后续AI处理，提高处理效率
            truncate_length = self.ai_parameters['content_extraction_length'] if self.ai_parameters['content_extraction_length'] < 2000 else len(extracted_content)
            content_for_ai = extracted_content[:truncate_length] if extracted_content else ""
            if len(extracted_content) > truncate_length:
                logging.info(f"内容已截取至前{truncate_length}字符用于AI处理（原长度: {len(extracted_content)} 字符）")
            
            # 第三步：生成100字摘要（使用截取后的内容）
            summary_start = time.time()
            summary = self._generate_content_summary(content_for_ai, file_name)
            summary_time = round(time.time() - summary_start, 3)
            timing_info['summary_generation_time'] = summary_time
            logging.info(f"内容摘要生成完成: {summary[:50]}...，耗时: {summary_time}秒")
            
            # 第四步：使用递归逐层匹配推荐最匹配的存放目录
            recommend_start = time.time()
            recommended_folder, level_tags, match_reason = self._recommend_target_folder_recursive(
                file_path, content_for_ai, summary, target_directory
            )
            recommend_time = round(time.time() - recommend_start, 3)
            timing_info['folder_recommendation_time'] = recommend_time
            
            total_time = round(time.time() - start_time, 3)
            timing_info['total_processing_time'] = total_time
            
            result = {
                'file_path': file_path,
                'file_name': file_name,
                'file_metadata': file_metadata,
                'extracted_content': extracted_content,
                'content_summary': summary,
                'recommended_folder': recommended_folder,
                'level_tags': level_tags,
                'match_reason': match_reason,
                'success': recommended_folder is not None,
                'timing_info': timing_info
            }
            
            if recommended_folder:
                logging.info(f"文件分析完成: {file_name} -> {recommended_folder}，总耗时: {total_time}秒")
                logging.info(f"摘要: {summary}")
                logging.info(f"推荐理由: {match_reason}")
                logging.info(f"各级标签: {level_tags}")
                logging.info(f"详细耗时 - 元数据提取: {metadata_time}秒, 内容提取: {extract_time}秒, 摘要生成: {summary_time}秒, 目录推荐: {recommend_time}秒")
            else:
                logging.warning(f"文件分析失败: {file_name}，总耗时: {total_time}秒")
            
            return result
            
        except Exception as e:
            total_time = round(time.time() - start_time, 3)
            timing_info['total_processing_time'] = total_time
            logging.error(f"文件分析失败: {e}，总耗时: {total_time}秒")
            return {
                'file_path': file_path,
                'file_name': Path(file_path).name,
                'file_metadata': {},
                'extracted_content': '',
                'content_summary': '',
                'recommended_folder': None,
                'level_tags': [],
                'match_reason': f"分析失败: {str(e)}",
                'success': False,
                'error': str(e),
                'timing_info': timing_info
            }
    
    def classify_file(self, file_path: str, target_directory: str) -> tuple:
        """
        保持原有的分类方法兼容性
        """
        result = self.analyze_and_classify_file(file_path, target_directory)
        return result['recommended_folder'], result['match_reason'], result['success']
    def _build_classification_prompt(self, file_path: str, target_directory: str) -> str:
        file_name = Path(file_path).name
        file_extension = Path(file_path).suffix.lower()
        file_content = ""
        content_readable = False
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                file_content = f.read()[:2000]
                printable_ratio = sum(1 for c in file_content if c.isprintable() or c.isspace()) / len(file_content) if file_content else 0
                if printable_ratio > 0.7:
                    content_readable = True
                else:
                    file_content = "文件内容为二进制格式，无法读取"
        except Exception:
            for encoding in ['gbk', 'gb2312', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        file_content = f.read()[:2000]
                        printable_ratio = sum(1 for c in file_content if c.isprintable() or c.isspace()) / len(file_content) if file_content else 0
                        if printable_ratio > 0.7:
                            content_readable = True
                            break
                        else:
                            file_content = "文件内容为二进制格式，无法读取"
                except Exception:
                    continue
            if not content_readable:
                file_content = "无法读取文件内容"
        directory_structure = self.get_directory_tree_structure(target_directory)
        prompt = f"""
你是一个专业的文件分类助手。请阅读原文件的前500个字，或者识别第一页内容，根据文件内容判断文件应该归类到下列哪个文件夹中，注意，每个文件根据文件内容只匹配一个文件夹，匹配原则是：

- 优先匹配：文件内容主题和文件夹名及文件路径名最相符
- 降级匹配：如果读取不到文件内容，则用原文件名和目标文件夹名及文件夹路径名来匹配

文件信息：
- 原文件名：{file_name}
- 文件扩展名：{file_extension}
- 文件内容：{file_content if content_readable else "无法读取内容"}

目标目录文件夹清单如下：
{directory_structure}

输出格式要求：
原文件名|目标文件夹|匹配理由

其中匹配理由格式为：
内容匹配：{{用不多于20个字描述匹配理由}}
或者
无法读取内容，采用文件夹名匹配：{{用不多于20个字描述匹配理由}}

请严格按照上述格式输出，只输出一行结果。
"""
        return prompt
    
    def _extract_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        提取文件元数据，包括创建时间、修改时间等
        
        Args:
            file_path: 文件路径
            
        Returns:
            包含文件元数据的字典
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise Exception(f"文件不存在: {file_path}")
            
            stat = file_path.stat()
            
            metadata = {
                'file_name': file_path.name,
                'file_extension': file_path.suffix.lower(),
                'file_size': stat.st_size,
                'created_time': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'full_path': str(file_path.absolute()),
                'relative_path': str(file_path)
            }
            
            return metadata
            
        except Exception as e:
            logging.error(f"提取文件元数据失败: {e}")
            return {
                'file_name': Path(file_path).name,
                'file_extension': Path(file_path).suffix.lower(),
                'file_size': 0,
                'created_time': datetime.now().isoformat(),
                'modified_time': datetime.now().isoformat(),
                'full_path': str(Path(file_path).absolute()),
                'relative_path': str(file_path),
                'error': str(e)
            }
    
    def _extract_file_content(self, file_path: str, max_length: int = 3000) -> str:
        """
        提取文件内容（从file_reader.py复制的完整实现）
        
        Args:
            file_path: 文件路径
            max_length: 最大提取长度
            
        Returns:
            提取的文件内容
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise Exception(f"文件不存在: {file_path}")
            
            file_extension = file_path.suffix.lower()
            logging.info(f"正在提取文件内容: {file_path.name} (类型: {file_extension})")
            
            # 根据文件类型选择不同的提取方法
            if file_extension == '.pdf':
                return self._extract_pdf_content(file_path, max_length)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_docx_content(file_path, max_length)
            elif file_extension in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv']:
                return self._extract_text_content(file_path, max_length)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
                return self._extract_image_info(file_path)
            else:
                # 尝试作为文本文件读取
                return self._extract_text_content(file_path, max_length)
                
        except Exception as e:
            logging.error(f"提取文件内容失败: {e}")
            return f"无法提取文件内容: {str(e)}"
    
    def _extract_pdf_content(self, file_path: Path, max_length: int) -> str:
        """
        提取PDF文件内容（从file_reader.py复制的完整实现）
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                content = ""
                
                # 读取前几页内容
                for page_num in range(min(3, len(pdf_reader.pages))):
                    page = pdf_reader.pages[page_num]
                    content += page.extract_text() + "\n"
                    
                    if len(content) >= max_length:
                        break
                
                return content[:max_length] if content else "PDF文件内容为空或无法提取"
                
        except Exception as e:
            return f"PDF文件读取失败: {str(e)}"
    
    def _extract_docx_content(self, file_path: Path, max_length: int) -> str:
        """
        提取Word文档内容（从file_reader.py复制的完整实现）
        """
        try:
            # 检查文件是否存在且可读
            if not file_path.exists():
                return "文件不存在"
            
            if not file_path.is_file():
                return "路径不是文件"
            
            # 检查文件大小
            file_size = file_path.stat().st_size
            if file_size == 0:
                return "文件为空"
            
            # 尝试读取Word文档
            try:
                doc = Document(file_path)
                content = ""
                
                # 提取段落内容
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # 跳过空段落
                        content += paragraph.text.strip() + "\n"
                        if len(content) >= max_length:
                            break
                
                # 如果段落内容为空，尝试提取表格内容
                if not content.strip():
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    content += cell.text.strip() + " "
                        content += "\n"
                        if len(content) >= max_length:
                            break
                
                return content[:max_length].strip() if content.strip() else "Word文档内容为空或无法提取"
                
            except Exception as docx_error:
                # 如果是.doc文件或损坏的.docx文件，尝试其他方法
                if file_path.suffix.lower() == '.doc':
                    return "不支持.doc格式，请转换为.docx格式"
                else:
                    return f"Word文档格式错误或文件损坏: {str(docx_error)}"
            
        except Exception as e:
            return f"Word文档读取失败: {str(e)}"
    
    def _extract_text_content(self, file_path: Path, max_length: int) -> str:
        """
        提取文本文件内容（从file_reader.py复制的完整实现）
        """
        try:
            # 尝试多种编码格式
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        content = f.read(max_length)
                        
                        # 检查内容是否可读
                        printable_ratio = sum(1 for c in content if c.isprintable() or c.isspace()) / len(content) if content else 0
                        if printable_ratio > 0.7:
                            return content
                        
                except Exception:
                    continue
            
            return "文件内容为二进制格式，无法读取"
            
        except Exception as e:
            return f"文本文件读取失败: {str(e)}"
    
    def _extract_image_info(self, file_path: Path) -> str:
        """
        提取图片文件信息（从file_reader.py复制的完整实现）
        """
        try:
            from PIL import Image
            # 兼容新版本Pillow
            try:
                from PIL import ImageDraw, ImageFont
            except ImportError:
                pass
            with Image.open(file_path) as img:
                info = f"图片文件信息:\n"
                info += f"格式: {img.format}\n"
                info += f"尺寸: {img.size[0]} x {img.size[1]}\n"
                info += f"模式: {img.mode}\n"
                
                # 如果有EXIF信息，提取一些基本信息
                if hasattr(img, '_getexif') and img._getexif():
                    info += "包含EXIF信息\n"
                
                return info
                
        except Exception as e:
            return f"图片文件信息提取失败: {str(e)}"
    
    def _generate_content_summary(self, content: str, file_name: str, summary_length: int = None) -> str:
        """
        生成文件内容摘要
        
        Args:
            content: 文件内容
            file_name: 文件名
            summary_length: 摘要长度（字数），如果为None则使用默认值
        """
        try:
            if not content or content.startswith("无法") or content.startswith("文件内容为二进制"):
                return f"无法生成摘要：{content[:50]}..."
            
            if not hasattr(self, 'ai_manager') or self.ai_manager is None:
                self.initialize_ollama()
            
            # 使用传入的摘要长度，如果没有则使用默认值
            target_length = summary_length if summary_length is not None else self.ai_parameters['summary_length']
            
            # 构建更明确的提示词，避免思考过程输出
            prompt = f"""请为以下文件内容生成一个{target_length}字以内的中文摘要。

文件名：{file_name}

文件内容：
{content}

要求：
1. 概括文件的主要内容和主题
2. 突出关键信息和要点
3. 语言简洁明了
4. 字数控制在{target_length}字以内
5. 直接输出摘要内容，不要包含任何思考过程或说明文字
6. 不要使用"<think>"标签或任何思考过程描述

摘要：/no_think"""

            # 在传递给大模型的完整内容最尾部添加/no_think标签
            final_prompt = prompt

            # 使用系统提示词来抑制思考过程
            messages = [
                {
                    'role': 'system',
                    'content': '你是一个专业的文档摘要助手。重要：不要输出任何推理过程、思考步骤或解释。直接按要求输出结果。只输出摘要内容，不要包含任何其他信息。'
                },
                {
                    'role': 'user',
                    'content': final_prompt
                }
            ]

            # 使用统一的AI管理器
            summary = chat_with_ai(messages)
            
            # 清理可能的思考过程标签和内容
            summary = summary.replace('<think>', '').replace('</think>', '').strip()
            
            # 移除常见的思考过程开头
            think_prefixes = [
                '好的，', '好，', '嗯，', '我来', '我需要', '首先，', '让我', '现在我要',
                '用户希望', '用户要求', '用户让我', '根据', '基于', '考虑到', '让我先仔细看看',
                '用户给了我这个查询', '用户给了我这个任务', '用户给了一个任务',
                '首先，我得看一下', '首先，我要理解', '首先，我得仔细看看',
                '好的，用户让我', '用户让我生成', '内容来自文件', '重点包括', '首先，我需要确认'
            ]
            
            # 更激进的清理：移除所有以思考过程开头的句子
            lines = summary.split('\n')
            cleaned_lines = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 检查是否以思考过程开头
                is_think_line = False
                for prefix in think_prefixes:
                    if line.lower().startswith(prefix.lower()):
                        is_think_line = True
                        break
                
                # 如果不是思考过程，保留这一行
                if not is_think_line:
                    cleaned_lines.append(line)
            
            # 如果清理后没有内容，尝试更简单的方法
            if not cleaned_lines:
                # 找到第一个不是思考过程的句子
                sentences = summary.split('。')
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                    
                    # 检查是否以思考过程开头
                    is_think_sentence = False
                    for prefix in think_prefixes:
                        if sentence.lower().startswith(prefix.lower()):
                            is_think_sentence = True
                            break
                    
                    if not is_think_sentence:
                        cleaned_lines.append(sentence)
                        break
            
            # 如果还是没有内容，使用原始摘要的最后一部分
            if not cleaned_lines:
                # 取最后100个字符作为摘要
                summary = summary[-100:] if len(summary) > 100 else summary
            
            # 重新组合清理后的内容
            if cleaned_lines:
                summary = '。'.join(cleaned_lines)
            
            # 确保摘要长度不超过限制
            if len(summary) > target_length:
                summary = summary[:target_length-3] + "..."
            
            return summary.strip()
            
        except Exception as e:
            logging.error(f"生成摘要失败: {e}")
            return f"摘要生成失败: {str(e)}"
    
    def _clean_ai_response(self, response: str) -> str:
        """清理AI响应中的思考过程"""
        if not response:
            return response
        
        print(f"🔧 开始清理AI响应，原始长度: {len(response)}")
        print(f"🔧 原始响应: {response}")
        
        # 清理可能的思考过程标签和内容
        response = response.replace('<think>', '').replace('</think>', '').strip()
        
        # 移除常见的思考过程开头
        think_prefixes = [
            '好的，', '好，', '嗯，', '我来', '我需要', '首先，', '让我', '现在我要',
            '用户希望', '用户要求', '用户让我', '根据', '基于', '考虑到', '让我先仔细看看',
            '用户给了我这个查询', '用户给了我这个任务', '用户给了一个任务',
            '首先，我得看一下', '首先，我要理解', '首先，我得仔细看看',
            '好的，用户让我', '用户让我生成', '内容来自文件', '重点包括', '首先，我需要确认'
        ]
        
        # 更激进的清理：移除所有以思考过程开头的句子
        lines = response.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检查是否以思考过程开头
            is_think_line = False
            for prefix in think_prefixes:
                if line.lower().startswith(prefix.lower()):
                    is_think_line = True
                    break
            
            # 如果不是思考过程，保留这一行
            if not is_think_line:
                cleaned_lines.append(line)
        
        print(f"🔧 清理后行数: {len(cleaned_lines)}")
        
        # 如果清理后没有内容，尝试更简单的方法
        if not cleaned_lines:
            print(f"🔧 清理后没有内容，尝试句子分割方法")
            # 找到第一个不是思考过程的句子
            sentences = response.split('。')
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                
                # 检查是否以思考过程开头
                is_think_sentence = False
                for prefix in think_prefixes:
                    if sentence.lower().startswith(prefix.lower()):
                        is_think_sentence = True
                        break
                
                if not is_think_sentence:
                    cleaned_lines.append(sentence)
                    break
        
        # 如果还是没有内容，使用原始响应的最后一部分
        if not cleaned_lines:
            print(f"🔧 仍然没有内容，使用最后100个字符")
            # 取最后100个字符
            response = response[-100:] if len(response) > 100 else response
            print(f"🔧 截断后响应: {response}")
            return response
        
        # 重新组合清理后的内容，保持换行符以支持多行推荐路径
        if cleaned_lines:
            response = '\n'.join(cleaned_lines)
        
        print(f"🔧 最终清理结果: {response}")
        return response.strip()
    
    def _get_level_directories(self, base_directory: str, level: int = 1) -> List[str]:
        """
        获取指定层级的所有目录
        
        Args:
            base_directory: 基础目录
            level: 层级（1表示一级目录，2表示二级目录等）
            
        Returns:
            该层级的所有目录名列表
        """
        try:
            base_path = Path(base_directory)
            if not base_path.exists():
                return []
            
            directories = []
            
            if level == 1:
                # 获取一级目录
                for item in base_path.iterdir():
                    if item.is_dir():
                        directories.append(item.name)
            else:
                # 获取指定层级目录
                for item in base_path.rglob('*'):
                    if item.is_dir():
                        # 计算相对路径的层级
                        relative_path = item.relative_to(base_path)
                        path_parts = relative_path.parts
                        if len(path_parts) == level:
                            directories.append(item.name)
            
            return directories
            
        except Exception as e:
            logging.error(f"获取{level}级目录失败: {e}")
            return []

    def _match_level_directory(self, file_info: Dict[str, Any], content: str, summary: str, 
                              base_directory: str, current_path: str, level: int) -> Tuple[str, str]:
        """
        匹配指定层级的目录
        
        Args:
            file_info: 文件信息字典
            content: 文件内容
            summary: 文件摘要
            base_directory: 基础目录
            current_path: 当前已匹配的路径
            level: 当前匹配的层级
            
        Returns:
            (匹配的目录名, 匹配理由)
        """
        try:
            # 获取当前层级的所有目录
            level_dirs = self._get_level_directories(base_directory, level)
            if not level_dirs:
                return "", "该层级没有子目录"
            
            file_name = file_info['file_name']
            file_extension = file_info['file_extension']
            
            # 获取用户自定义分类规则
            custom_rules = self.classification_rules_manager.get_rules_for_prompt(level_dirs)
            
            # 构建匹配提示词
            prompt = f"""你是一个专业的文件分类专家。请根据文件信息在当前层级目录中选择最匹配的一个。

文件信息：
- 文件名：{file_name}
- 文件扩展名：{file_extension}
- 内容摘要：{summary[:200] if summary else '无摘要'}
- 当前路径：{current_path if current_path else '根目录'}

当前层级可选目录（必须严格从以下列表中选择一个）：
{chr(10).join(f"{i+1}. {dir_name}" for i, dir_name in enumerate(level_dirs))}

{custom_rules}

匹配优先级：
1. 最高优先级：目录名是时间命名（如年份、月份），而文件名中包含对应时间
2. 高优先级：目录名直接包含在文件名中
3. 中优先级：目录名与文件内容主题高度相关
4. 低优先级：根据文件类型和扩展名匹配

请只返回一个最匹配的目录名，不要包含任何其他内容："""

            # 调用AI进行匹配
            messages = [
                {
                    'role': 'system',
                    'content': '你是一个专业的文件分类专家。重要：不要输出任何推理过程、思考步骤或解释。直接按要求输出结果。只输出目录名，不要包含任何其他信息。'
                },
                {
                    'role': 'user',
                    'content': prompt
                }
            ]
            
            result = chat_with_ai(messages)
            result = result.strip()
            
            # 清理结果
            if any(keyword in result.lower() for keyword in ['好，', '嗯，', '我来', '我需要', '首先，', '让我']):
                sentences = result.split('。')
                if len(sentences) > 1:
                    result = '。'.join(sentences[1:]).strip()
            
            # 验证结果是否在可选目录中
            if result in level_dirs:
                # 确定匹配理由
                match_reason = self._determine_match_reason(file_name, result, summary)
                return result, match_reason
            else:
                # 如果AI返回的结果不在列表中，尝试模糊匹配
                for dir_name in level_dirs:
                    if self._fuzzy_match(file_name, dir_name, summary):
                        match_reason = f"模糊匹配到: {dir_name}"
                        return dir_name, match_reason
                
                # 如果都没有匹配到，返回第一个目录
                return level_dirs[0], f"默认选择第一个目录: {level_dirs[0]}"
                
        except Exception as e:
            logging.error(f"匹配{level}级目录失败: {e}")
            return "", f"匹配失败: {str(e)}"

    def _determine_match_reason(self, file_name: str, dir_name: str, summary: str) -> str:
        """
        确定匹配理由
        
        Args:
            file_name: 文件名
            dir_name: 目录名
            summary: 文件摘要
            
        Returns:
            匹配理由
        """
        # 检查时间匹配
        if self._is_time_match(file_name, dir_name):
            return f"时间匹配: 文件名包含时间信息，匹配到时间目录 {dir_name}"
        
        # 检查文件名包含目录名
        if dir_name.lower() in file_name.lower():
            return f"文件名匹配: 文件名包含目录名 {dir_name}"
        
        # 检查内容主题匹配
        if summary and any(keyword in summary.lower() for keyword in dir_name.lower().split()):
            return f"内容主题匹配: 文件内容与目录 {dir_name} 主题相关"
        
        return f"AI智能分类: 根据文件内容匹配到 {dir_name}"

    def _is_time_match(self, file_name: str, dir_name: str) -> bool:
        """
        检查是否为时间匹配
        
        Args:
            file_name: 文件名
            dir_name: 目录名
            
        Returns:
            是否为时间匹配
        """
        # 提取文件名中的年份
        year_pattern = r'\b(19|20)\d{2}\b'
        file_years = re.findall(year_pattern, file_name)
        
        # 检查目录名是否包含年份
        dir_years = re.findall(year_pattern, dir_name)
        
        if file_years and dir_years:
            return any(fy in dy for fy in file_years for dy in dir_years)
        
        return False

    def _fuzzy_match(self, file_name: str, dir_name: str, summary: str) -> bool:
        """
        模糊匹配
        
        Args:
            file_name: 文件名
            dir_name: 目录名
            summary: 文件摘要
            
        Returns:
            是否模糊匹配
        """
        # 简单的关键词匹配
        file_keywords = set(file_name.lower().split())
        dir_keywords = set(dir_name.lower().split())
        
        # 检查是否有共同关键词
        common_keywords = file_keywords.intersection(dir_keywords)
        if len(common_keywords) > 0:
            return True
        
        # 检查摘要中的关键词
        if summary:
            summary_keywords = set(summary.lower().split())
            summary_dir_common = summary_keywords.intersection(dir_keywords)
            if len(summary_dir_common) > 0:
                return True
        
        return False

    def _recommend_target_folder_recursive(self, file_path: str, content: str, summary: str, 
                                         target_directory: str, retry_count: int = 0) -> Tuple[str, List[str], str]:
        """
        递归逐层匹配目标文件夹
        
        Args:
            file_path: 文件路径
            content: 文件内容
            summary: 文件摘要
            target_directory: 目标目录
            retry_count: 重试次数
            
        Returns:
            (完整匹配路径, 各级标签列表, 匹配理由)
        """
        try:
            file_info = self.file_cache.get(file_path, {})
            if not file_info:
                file_info = self._extract_file_metadata(file_path)
                self.file_cache[file_path] = file_info
            
            current_path = ""
            full_path = ""
            level_tags = []
            match_reasons = []
            
            level = 1
            max_levels = 10  # 防止无限递归
            
            while level <= max_levels:
                # 构建当前层级的完整路径
                if current_path:
                    current_full_path = os.path.join(target_directory, current_path)
                else:
                    current_full_path = target_directory
                
                # 匹配当前层级
                matched_dir, match_reason = self._match_level_directory(
                    file_info, content, summary, current_full_path, current_path, level
                )
                
                if not matched_dir:
                    break
                
                # 更新路径和标签
                if current_path:
                    current_path = os.path.join(current_path, matched_dir)
                else:
                    current_path = matched_dir
                
                full_path = current_path
                level_tags.append(matched_dir)
                match_reasons.append(f"第{level}级: {match_reason}")
                
                # 检查下一级是否有子目录
                next_level_path = os.path.join(current_full_path, matched_dir)
                next_level_dirs = self._get_level_directories(next_level_path, 1)
                
                if not next_level_dirs:
                    break
                
                level += 1
            
            # 缓存标签信息
            self.level_tags[file_path] = level_tags
            
            # 合并匹配理由
            combined_reason = "; ".join(match_reasons)
            
            return full_path, level_tags, combined_reason
            
        except Exception as e:
            logging.error(f"递归匹配目标文件夹失败: {e}")
            return "", [], f"匹配失败: {str(e)}"

    def _clear_file_cache(self, file_path: str) -> None:
        """
        清理指定文件的缓存，但保留全局缓存
        
        Args:
            file_path: 要清理缓存的文件路径
        """
        try:
            # 清理文件缓存
            if file_path in self.file_cache:
                del self.file_cache[file_path]
            
            # 清理标签缓存
            if file_path in self.level_tags:
                del self.level_tags[file_path]
                
            logging.debug(f"已清理文件缓存: {file_path}")
            
        except Exception as e:
            logging.error(f"清理文件缓存失败: {e}")

    def _recommend_target_folder(self, file_path: str, content: str, summary: str, target_directory: str, retry_count: int = 0) -> tuple:
        """
        基于文件内容和摘要推荐最匹配的目标文件夹
        使用新的递归逐层匹配逻辑
        """
        try:
            file_name = Path(file_path).name
            file_full_path = str(Path(file_path).absolute())
            
            print(f"📄 源文件完整路径: {file_full_path}")
            if retry_count > 0:
                print(f"🔄 第 {retry_count} 次重试AI分类")
            
            # 存储源文件完整路径
            self.source_file_path = file_full_path
            
            # 使用新的递归匹配逻辑
            recommended_folder, level_tags, match_reason = self._recommend_target_folder_recursive(
                file_path, content, summary, target_directory, retry_count
            )
            
            # 检查分类质量：如果推荐的是过于宽泛的分类，尝试重新分类
            if recommended_folder and retry_count < 2:
                # 检查路径深度，如果太浅可能是过于宽泛的分类
                path_depth = len([p for p in recommended_folder.split('\\') + recommended_folder.split('/') if p.strip()])
                if path_depth <= 1:  # 路径深度小于等于1可能是过于宽泛
                    print(f"⚠️  AI推荐了过于宽泛的分类: {recommended_folder}，准备第 {retry_count + 1} 次重试...")
                    logging.warning(f"AI推荐了过于宽泛的分类: {recommended_folder}，准备第 {retry_count + 1} 次重试")
                    
                    # 在重试时强调要选择更具体的分类
                    return self._recommend_target_folder(file_path, content, summary, target_directory, retry_count + 1)
            
            # 检查是否匹配失败，如果是且未超过重试次数，则重试
            if not recommended_folder and retry_count < 2:  # 最多重试2次
                print(f"⚠️  AI分类失败，准备第 {retry_count + 1} 次重试...")
                logging.warning(f"AI分类失败，准备第 {retry_count + 1} 次重试")
                
                # 递归调用自身进行重试
                return self._recommend_target_folder(file_path, content, summary, target_directory, retry_count + 1)
            
            return recommended_folder, match_reason
            
        except Exception as e:
            logging.error(f"推荐目标文件夹失败: {e}")
            return None, f"推荐失败: {str(e)}"
    def _extract_recommended_paths(self, ai_result: str) -> List[str]:
        """
        从AI返回结果中提取三个推荐路径
        
        Args:
            ai_result: AI返回的原始结果
            
        Returns:
            List[str]: 提取到的推荐路径列表
        """
        recommended_paths = []
        
        try:
            lines = ai_result.strip().split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # 查找包含推荐路径的行
                if any(keyword in line for keyword in ['第一推荐：', '第二推荐：', '第三推荐：', '推荐：']):
                    # 提取冒号后的路径
                    if '：' in line:
                        path = line.split('：', 1)[1].strip()
                        if path and path not in recommended_paths:
                            recommended_paths.append(path)
                elif line.startswith('【') and '】' in line:
                    # 直接识别以【开头的路径格式
                    if line not in recommended_paths:
                        recommended_paths.append(line)
            
            # 如果没有找到格式化的推荐，尝试按行提取前三个有效路径
            if not recommended_paths:
                for line in lines:
                    line = line.strip()
                    if line and line.startswith('【') and '】' in line:
                        if line not in recommended_paths:
                            recommended_paths.append(line)
                        if len(recommended_paths) >= 3:
                            break
            
            # 限制最多返回3个路径
            return recommended_paths[:3]
            
        except Exception as e:
            logging.error(f"提取推荐路径时出错: {e}")
            return []
    
    def _parse_classification_result(self, result: str, target_directory: str, file_content: str = None, file_summary: str = None) -> tuple:
        """
        解析AI分类结果，支持相关性评分排序和三个推荐路径的依次验证
        
        Args:
            result: AI返回的分类结果
            target_directory: 目标目录路径
            file_content: 文件内容（用于相关性计算）
            file_summary: 文件摘要（用于相关性计算）
            
        Returns:
            tuple: (推荐文件夹路径, 匹配理由)
        """
        try:
            result = result.strip()
            logging.info(f"正在解析AI分类结果: {result[:100]}...")
            
            print(f"🔍 解析AI分类结果: {result}")
            
            # 清理AI返回结果中的思考过程标签和内容
            result_clean = self._clean_ai_response(result)
            logging.info(f"清理后的分类结果: {result_clean[:100]}...")
            print(f"🧹 清理后的结果: {result_clean}")
            
            # 提取三个推荐路径
            recommended_paths = self._extract_recommended_paths(result_clean)
            
            if not recommended_paths:
                logging.warning(f"无法从AI结果中提取推荐路径")
                print(f"⚠️ 无法提取推荐路径")
                return None, f"AI智能分类失败，无法解析推荐路径: {result_clean[:50]}..."
            
            print(f"📋 提取到 {len(recommended_paths)} 个推荐路径: {recommended_paths}")
            
            # 如果有文件内容和摘要，进行相关性分析和排序
            if file_content and file_summary:
                print(f"🔍 开始相关性分析和排序...")
                sorted_paths = self._filter_irrelevant_folders(file_content, file_summary, recommended_paths, target_directory)
                
                if sorted_paths:
                    recommended_paths = sorted_paths
                    print(f"✅ 相关性排序完成，按评分排序后的路径: {sorted_paths}")
                else:
                    print(f"⚠️ 相关性排序后无相关路径，使用原始推荐")
            
            # 依次验证每个推荐路径（现在按相关性评分排序）
            for i, recommended_path in enumerate(recommended_paths, 1):
                target_path = Path(target_directory) / recommended_path
                
                print(f"🔍 验证第{i}个推荐路径: {recommended_path}")
                print(f"🎯 完整目标路径: {target_path}")
                
                # 验证路径是否存在且为目录
                if target_path.exists() and target_path.is_dir():
                    logging.info(f"找到有效的目标文件夹(第{i}选择): {recommended_path}")
                    print(f"✅ 找到有效目标文件夹(第{i}选择): {recommended_path}")
                    
                    # 存储推荐的完整路径
                    self.recommended_folder_path = recommended_path
                    print(f"💾 存储推荐路径: {self.recommended_folder_path}")
                    
                    # 生成有意义的匹配理由
                    match_reason = f"AI智能分类(第{i}选择)：根据文件内容匹配到 {recommended_path}"
                    
                    return recommended_path, match_reason
                else:
                    logging.warning(f"第{i}个推荐路径不存在: {target_path}")
                    print(f"❌ 第{i}个推荐路径不存在: {recommended_path}")
            
            # 如果所有推荐路径都无效
            logging.error(f"所有推荐路径都无效: {recommended_paths}")
            print(f"💥 所有推荐路径都无效")
            
            return None, f"AI智能分类失败，三个推荐路径都不存在: {', '.join(recommended_paths[:3])}"
            
        except Exception as e:
            logging.error(f"解析分类结果失败: {e}, 原始结果: {result}")
            return None, f"解析分类结果时发生错误: {str(e)}"
    def scan_source_files(self, source_directory: str) -> List[str]:
        try:
            source_path = Path(source_directory)
            if not source_path.exists():
                raise FileOrganizerError(f"源目录不存在: {source_directory}")
            files = []
            for item in source_path.rglob('*'):
                if item.is_file():
                    files.append(str(item))
            logging.info(f"扫描到 {len(files)} 个待整理文件")
            return files
        except Exception as e:
            raise FileOrganizerError(f"扫描源文件失败: {e}")
    def scan_files(self, source_directory: str) -> List[Dict[str, object]]:
        try:
            source_path = Path(source_directory)
            if not source_path.exists():
                raise FileOrganizerError(f"源目录不存在: {source_directory}")
            files = []
            for item in source_path.rglob('*'):
                if item.is_file():
                    files.append({
                        'name': item.name,
                        'path': str(item),
                        'size': item.stat().st_size,
                        'extension': item.suffix.lower()
                    })
            logging.info(f"扫描到 {len(files)} 个待整理文件")
            return files
        except Exception as e:
            raise FileOrganizerError(f"扫描源文件失败: {e}")

    def preview_classification(self, source_directory: str, target_directory: str) -> List[Dict[str, object]]:
        try:
            source_files = self.scan_source_files(source_directory)
            if not source_files:
                raise FileOrganizerError("源目录中没有找到文件")
            preview_results = []
            logging.info(f"开始预览分类，共 {len(source_files)} 个文件")
            preview_count = len(source_files)
            for i, file_path in enumerate(source_files, 1):
                file_name = Path(file_path).name
                try:
                    logging.info(f"正在处理文件 {i}/{preview_count}: {file_name}")
                    print(f"\n=== 处理文件 {i}/{preview_count}: {file_name} ===")
                    # 使用新的分析方法
                    analysis_result = self.analyze_and_classify_file(file_path, target_directory)
                    
                    if analysis_result['success']:
                        timing_info = analysis_result.get('timing_info', {})
                        preview_results.append({
                            'file_path': file_path,
                            'file_name': file_name,
                            'target_folder': analysis_result['recommended_folder'],
                            'match_reason': analysis_result['match_reason'],
                            'classification_method': 'AI_Enhanced',
                            'success': True,
                            'extracted_content': analysis_result['extracted_content'][:200] + "..." if len(analysis_result['extracted_content']) > 200 else analysis_result['extracted_content'],
                            'content_summary': analysis_result['content_summary'],
                            'timing_info': timing_info
                        })
                        total_time = timing_info.get('total_processing_time', 0)
                        extract_time = timing_info.get('content_extraction_time', 0)
                        summary_time = timing_info.get('summary_generation_time', 0)
                        recommend_time = timing_info.get('folder_recommendation_time', 0)
                        
                        logging.info(f"文件 {file_name} 分析成功: {analysis_result['recommended_folder']} ({analysis_result['match_reason']})，总耗时: {total_time}秒")
                        logging.info(f"内容摘要: {analysis_result['content_summary']}")
                        logging.info(f"详细耗时 - 内容提取: {extract_time}秒, 摘要生成: {summary_time}秒, 目录推荐: {recommend_time}秒")
                    else:
                        timing_info = analysis_result.get('timing_info', {})
                        preview_results.append({
                            'file_path': file_path,
                            'file_name': file_name,
                            'target_folder': None,
                            'match_reason': analysis_result['match_reason'],
                            'classification_method': 'Failed',
                            'success': False,
                            'error': analysis_result.get('error', analysis_result['match_reason']),
                            'extracted_content': analysis_result.get('extracted_content', ''),
                            'content_summary': analysis_result.get('content_summary', ''),
                            'timing_info': timing_info
                        })
                        total_time = timing_info.get('total_processing_time', 0)
                        logging.warning(f"文件 {file_name} 分析失败: {analysis_result['match_reason']}，总耗时: {total_time}秒")
                except Exception as e:
                    error_msg = f"处理文件时出错: {str(e)}"
                    logging.error(f"文件 {file_name} 处理异常: {e}")
                    preview_results.append({
                        'file_path': file_path,
                        'file_name': file_name,
                        'target_folder': None,
                        'match_reason': error_msg,
                        'classification_method': 'Error',
                        'success': False,
                        'error': error_msg
                    })
            success_count = sum(1 for result in preview_results if result['success'])
            logging.info(f"预览分类完成，成功分类 {success_count}/{len(preview_results)} 个文件")
            
            # 生成预览结果JSON文件（格式与preview_ai_result.json保持一致）
            try:
                ai_preview_results = []
                for result in preview_results:
                    if result['success']:
                        ai_preview_result = {
                            "源文件路径": result['file_path'],
                            "文件摘要": result['content_summary'],
                            "最匹配的目标目录": result['target_folder'],
                            "匹配理由": result['match_reason'],
                            "处理耗时信息": {
                                "总耗时(秒)": result['timing_info'].get('total_processing_time', 0),
                                "内容提取耗时(秒)": result['timing_info'].get('content_extraction_time', 0),
                                "摘要生成耗时(秒)": result['timing_info'].get('summary_generation_time', 0),
                                "目录推荐耗时(秒)": result['timing_info'].get('folder_recommendation_time', 0)
                            }
                        }
                        ai_preview_results.append(ai_preview_result)
                
                # 保存预览结果到JSON文件
                preview_result_file = 'preview_ai_result.json'
                with open(preview_result_file, 'w', encoding='utf-8') as f:
                    json.dump(ai_preview_results, f, ensure_ascii=False, indent=2)
                logging.info(f"AI预览结果已保存到: {preview_result_file}")
                
            except Exception as e:
                logging.warning(f"生成预览结果文件失败: {e}")
            
            return preview_results
        except Exception as e:
            raise FileOrganizerError(f"预览分类失败: {e}")
    def organize_file(self, file_path: str, target_directory: str) -> Tuple[bool, str]:
        try:
            if not hasattr(self, 'ai_manager') or self.ai_manager is None:
                self.initialize_ollama()
            
            file_path_obj = Path(file_path)
            filename = file_path_obj.name
            source_file_full_path = str(file_path_obj.absolute())
            
            print(f"📄 源文件完整路径: {source_file_full_path}")
            
            # 使用AI分类获取推荐文件夹
            target_folder, match_reason, success = self.classify_file(str(file_path_obj), target_directory)
            if not success or not target_folder:
                return False, "无法确定目标文件夹"
            
            # 使用存储的推荐路径变量
            if hasattr(self, 'recommended_folder_path') and self.recommended_folder_path:
                target_folder_full_path = Path(target_directory) / self.recommended_folder_path
                print(f"🔧 使用存储的推荐路径: {self.recommended_folder_path}")
                print(f"🎯 目标文件夹完整路径: {target_folder_full_path}")
            else:
                target_folder_full_path = Path(target_directory) / target_folder
                print(f"⚠️  使用默认路径: {target_folder}")
                print(f"🎯 目标文件夹完整路径: {target_folder_full_path}")
            
            # 构建完整的迁移命令信息
            migration_info = {
                'source_file': source_file_full_path,
                'target_folder': str(target_folder_full_path),
                'filename': filename,
                'match_reason': match_reason
            }
            
            print(f"📋 迁移信息: {migration_info}")
            
            # 检查并创建年份子文件夹
            final_target_folder = self._check_and_create_year_folder(target_folder_full_path, filename)
            
            # 确保目标文件夹存在
            final_target_folder.mkdir(parents=True, exist_ok=True)
            
            # 构建目标文件路径
            target_file_path = final_target_folder / filename
            if target_file_path.exists():
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                name_parts = filename.rsplit('.', 1)
                if len(name_parts) == 2:
                    new_filename = f"{name_parts[0]}_{timestamp}.{name_parts[1]}"
                else:
                    new_filename = f"{filename}_{timestamp}"
                target_file_path = final_target_folder / new_filename
                print(f"⚠️  文件名冲突，重命名为: {new_filename}")
            
            # 执行文件复制
            shutil.copy2(source_file_full_path, str(target_file_path))
            
            # 记录迁移日志
            migration_log = f"文件迁移完成: {source_file_full_path} -> {target_file_path}"
            logging.info(migration_log)
            print(f"✅ {migration_log}")
            
            return True, str(target_file_path)
        except Exception as e:
            error_msg = f"整理文件失败: {e}"
            logging.error(error_msg)
            return False, error_msg
    def organize_files(self, files=None, target_base_dir=None, copy_mode=True, source_directory=None, target_directory=None, dry_run=False, progress_callback=None) -> Dict[str, object]:
        try:
            if source_directory and target_directory:
                files = self.scan_files(source_directory)
                target_base_dir = target_directory
                copy_mode = True
            if not files or not target_base_dir:
                raise FileOrganizerError("缺少必要参数")
            log_session_name = None
            if self.enable_transfer_log and self.transfer_log_manager and not dry_run:
                try:
                    session_name = f"organize_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                    log_session_name = self.transfer_log_manager.start_transfer_session(session_name)
                    logging.info(f"开始转移日志会话: {session_name}")
                except Exception as e:
                    logging.warning(f"启动转移日志会话失败: {e}")
            if not files:
                raise FileOrganizerError("没有文件需要整理")
            results = {
                'total_files': len(files),
                'processed_files': 0,
                'successful_moves': 0,
                'failed_moves': 0,
                'skipped_files': 0,
                'success': [],
                'failed': [],
                'errors': [],
                'move_details': [],
                'ai_responses': [],
                'start_time': datetime.now(),
                'end_time': None,
                'transfer_log_file': log_session_name,
                'dry_run': dry_run
            }
            
            # 初始化AI结果文件
            ai_result_file = 'ai_organize_result.json'
            # 不再在内存中保存完整的ai_results列表，只保存必要的迁移信息
            migration_queue = []  # 只保存源路径和目标路径的简单信息
            
            # 确保AI结果文件存在，但不清空现有内容
            if not os.path.exists(ai_result_file):
                with open(ai_result_file, 'w', encoding='utf-8') as f:
                    json.dump([], f, ensure_ascii=False, indent=2)
                logging.info(f"创建新的AI结果文件: {ai_result_file}")
            else:
                logging.info(f"AI结果文件已存在，将追加新记录: {ai_result_file}")
            logging.info(f"开始安全文件整理，共 {len(files)} 个文件")
            print(f"\n=== 开始AI智能文件整理 ===")
            print(f"源目录: {source_directory if source_directory else '指定文件列表'}")
            print(f"目标目录: {target_base_dir}")
            print(f"待处理文件总数: {len(files)}")
            print(f"操作模式: {'复制' if copy_mode else '移动'}")
            print("=" * 50)
            
            for i, file_info in enumerate(files, 1):
                file_path = str(file_info['path'])
                filename = str(file_info['name'])
                
                # 调用进度回调
                if progress_callback:
                    progress_callback(i, len(files), filename)
                
                # 控制台输出当前处理进度
                print(f"\n[{i}/{len(files)}] 正在处理: {filename}")
                
                try:
                    logging.info(f"正在处理文件 {i}/{len(files)}: {filename}")
                    print(f"  🔍 正在分析文件内容...", end="", flush=True)
                    
                    # 获取详细的分析结果（只调用一次）
                    analysis_result = self.analyze_and_classify_file(file_path, target_base_dir)
                    target_folder = analysis_result.get('recommended_folder')
                    match_reason = analysis_result.get('match_reason', '')
                    success = analysis_result.get('success', False)
                    
                    results['ai_responses'].append({
                        'file_name': filename,
                        'target_folder': target_folder,
                        'match_reason': match_reason,
                        'success': success
                    })
                    
                    # 构建AI结果项（基础信息）
                    level_tags = analysis_result.get('level_tags', [])
                    file_metadata = analysis_result.get('file_metadata', {})
                    
                    ai_result_item = {
                        "处理时间": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        "文件名": filename,
                        "源文件路径": file_path,
                        "文件摘要": analysis_result.get('content_summary', ''),
                        "最匹配的目标目录": analysis_result.get('recommended_folder', ''),
                        "匹配理由": analysis_result.get('match_reason', ''),
                        "处理耗时信息": {
                            "总耗时(秒)": analysis_result.get('timing_info', {}).get('total_processing_time', 0),
                            "元数据提取耗时(秒)": analysis_result.get('timing_info', {}).get('metadata_extraction_time', 0),
                            "内容提取耗时(秒)": analysis_result.get('timing_info', {}).get('content_extraction_time', 0),
                            "摘要生成耗时(秒)": analysis_result.get('timing_info', {}).get('summary_generation_time', 0),
                            "目录推荐耗时(秒)": analysis_result.get('timing_info', {}).get('folder_recommendation_time', 0)
                        },
                        "文件元数据": {
                            "创建时间": file_metadata.get('created_time', ''),
                            "修改时间": file_metadata.get('modified_time', ''),
                            "文件大小": file_metadata.get('file_size', 0)
                        }
                    }
                    
                    # 添加各级标签
                    for i, tag in enumerate(level_tags, 1):
                        ai_result_item[f"{i}级标签"] = tag
                    
                    # 保存到迁移队列，等待迁移成功后写入完整信息
                    migration_queue.append({
                        'source_path': file_path,
                        'target_folder': target_folder,
                        'filename': filename,
                        'match_reason': match_reason,
                        'ai_result_item': ai_result_item
                    })
                    
                    if not success or not target_folder:
                        error_msg = f"文件 {filename} 分类失败: {match_reason}，已跳过，未做任何处理"
                        print(f"\r  ❌ 分类失败: {match_reason}")
                        logging.warning(error_msg)
                        results['errors'].append(error_msg)
                        results['skipped_files'] += 1
                        results['failed'].append({
                            'source_path': file_path,
                            'error': error_msg
                        })
                        
                        # 更新迁移队列中最后一项的失败状态并立即写入
                        if migration_queue:
                            migration_queue[-1]['ai_result_item'].update({
                                "处理状态": "分类失败",
                                "错误信息": match_reason
                            })
                            # 立即写入失败记录
                            self._append_ai_result_to_file(ai_result_file, migration_queue[-1]['ai_result_item'])
                        
                        # 清理当前文件的缓存
                        self._clear_file_cache(file_path)
                        
                        continue
                    
                    print(f"\r  ✅ 推荐目录: {target_folder}")
                    print(f"     理由: {match_reason}")
                    target_folder_path = Path(target_base_dir) / target_folder
                    if not target_folder_path.exists():
                        error_msg = f"目标文件夹不存在: {target_folder}"
                        logging.error(error_msg)
                        results['errors'].append(error_msg)
                        results['failed_moves'] += 1
                        results['failed'].append({
                            'source_path': file_path,
                            'error': error_msg
                        })
                        continue
                    original_target_path = target_folder_path / filename
                    target_file_path = original_target_path
                    if target_file_path.exists():
                        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                        name_parts = filename.rsplit('.', 1)
                        if len(name_parts) == 2:
                            new_filename = f"{name_parts[0]}_{timestamp}.{name_parts[1]}"
                        else:
                            new_filename = f"{filename}_{timestamp}"
                        target_file_path = target_folder_path / new_filename
                        logging.info(f"文件名冲突，重命名为: {target_file_path.name}")
                    if not dry_run:
                        try:
                            if not Path(file_path).exists():
                                error_msg = f"源文件不存在: {file_path}"
                                logging.error(error_msg)
                                results['errors'].append(error_msg)
                                results['failed_moves'] += 1
                                results['failed'].append({
                                    'source_path': file_path,
                                    'error': error_msg
                                })
                                continue
                            if copy_mode:
                                shutil.copy2(file_path, str(target_file_path))
                                operation = "copy"
                                operation_cn = "复制"
                            else:
                                shutil.move(file_path, str(target_file_path))
                                operation = "move"
                                operation_cn = "移动"
                            if target_file_path.exists():
                                if not copy_mode and Path(file_path).exists():
                                    error_msg = f"文件移动验证失败: {filename}"
                                    logging.error(error_msg)
                                    results['errors'].append(error_msg)
                                    results['failed_moves'] += 1
                                    results['failed'].append({
                                        'source_path': file_path,
                                        'error': error_msg
                                    })
                                    continue
                                print(f"     ✅ {operation_cn}成功: {filename} -> {target_folder}")
                                logging.info(f"文件安全{operation_cn}成功: {filename} -> {target_folder} ({match_reason})")
                                results['successful_moves'] += 1
                                
                                # 更新迁移队列中对应项的最终路径并立即写入
                                for queue_item in migration_queue:
                                    if queue_item['source_path'] == file_path:
                                        queue_item['ai_result_item'].update({
                                            "最终目标路径": str(target_file_path),
                                            "操作类型": operation_cn,
                                            "处理状态": "成功"
                                        })
                                        # 立即写入成功记录
                                        self._append_ai_result_to_file(ai_result_file, queue_item['ai_result_item'])
                                        break
                                
                                # 清理当前文件的缓存
                                self._clear_file_cache(file_path)
                            else:
                                error_msg = f"文件{operation_cn}验证失败: {filename}"
                                logging.error(error_msg)
                                results['errors'].append(error_msg)
                                results['failed_moves'] += 1
                                results['failed'].append({
                                    'source_path': file_path,
                                    'error': error_msg
                                })
                                continue
                        except Exception as move_error:
                            error_msg = f"{operation_cn}文件 {filename} 时出错: {move_error}"
                            logging.error(error_msg)
                            results['errors'].append(error_msg)
                            results['failed_moves'] += 1
                            results['failed'].append({
                                'source_path': file_path,
                                'error': error_msg
                            })
                            continue
                    else:
                        operation = "copy" if copy_mode else "move"
                        operation_cn = "复制" if copy_mode else "移动"
                        logging.info(f"[试运行] 文件将{operation_cn}: {filename} -> {target_folder} ({match_reason})")
                        results['successful_moves'] += 1
                        
                        # 试运行模式下更新迁移队列中对应项并立即写入
                        for queue_item in migration_queue:
                            if queue_item['source_path'] == file_path:
                                queue_item['ai_result_item'].update({
                                    "最终目标路径": str(target_file_path),
                                    "操作类型": operation_cn,
                                    "处理状态": "试运行成功"
                                })
                                # 立即写入试运行记录
                                self._append_ai_result_to_file(ai_result_file, queue_item['ai_result_item'])
                                break
                        
                        # 清理当前文件的缓存
                        self._clear_file_cache(file_path)
                    if self.enable_transfer_log and self.transfer_log_manager and not dry_run:
                        try:
                            file_size_raw = file_info.get('size', 0)
                            file_size = int(file_size_raw) if isinstance(file_size_raw, (int, float, str)) and str(file_size_raw).isdigit() else 0
                            self.transfer_log_manager.log_transfer_operation(
                                source_path=file_path,
                                target_path=str(target_file_path),
                                operation_type=operation,
                                target_folder=target_folder,
                                success=True,
                                file_size=file_size
                            )
                        except Exception as e:
                            logging.warning(f"记录转移日志失败: {e}")
                    results['success'].append({
                        'source_path': file_path,
                        'target_path': str(target_file_path),
                        'target_folder': target_folder,
                        'operation': operation
                    })
                    results['move_details'].append({
                        'source_file': file_path,
                        'file_name': filename,
                        'target_folder': target_folder,
                        'target_path': str(target_file_path),
                        'match_reason': match_reason,
                        'classification_method': 'AI',
                        'renamed': str(target_file_path) != str(original_target_path),
                        'operation': operation
                    })
                except Exception as e:
                    error_msg = f"处理文件 {filename} 时出现异常: {e}"
                    logging.error(error_msg)
                    results['errors'].append(error_msg)
                    results['failed_moves'] += 1
                    if self.enable_transfer_log and self.transfer_log_manager and not dry_run:
                        try:
                            operation_type = "copy" if copy_mode else "move"
                            self.transfer_log_manager.log_transfer_operation(
                                source_path=file_path,
                                target_path="",
                                operation_type=operation_type,
                                target_folder="",
                                success=False,
                                error_message=error_msg
                            )
                        except Exception as log_e:
                            logging.warning(f"记录失败转移日志失败: {log_e}")
                    results['failed'].append({
                        'source_path': file_path,
                        'error': error_msg
                    })
                finally:
                    results['processed_files'] += 1
            results['end_time'] = datetime.now()
            
            # 输出处理完成总结
            print(f"\n=== 文件整理完成 ===")
            print(f"总文件数: {results['total_files']}")
            print(f"成功处理: {results['successful_moves']} 个")
            print(f"处理失败: {results['failed_moves']} 个")
            print(f"跳过文件: {results['skipped_files']} 个")
            duration = (results['end_time'] - results['start_time']).total_seconds()
            print(f"总耗时: {duration:.1f} 秒")
            print("=" * 50)
            
            # AI结果已在处理过程中实时写入
            results['ai_result_file'] = ai_result_file
            logging.info(f"AI分析结果已实时写入: {ai_result_file}")
            
            if self.enable_transfer_log and self.transfer_log_manager and not dry_run:
                try:
                    session_summary = self.transfer_log_manager.end_transfer_session()
                    logging.info(f"转移日志会话结束: {session_summary}")
                except Exception as e:
                    logging.warning(f"结束转移日志会话失败: {e}")
            
            logging.info(f"安全文件整理完成: 成功 {results['successful_moves']}, 失败 {results['failed_moves']}, 跳过 {results['skipped_files']}")
            
            # 删除源文件功能已移至GUI界面，不再在控制台询问
            if not dry_run and results['successful_moves'] > 0:
                logging.info(f"成功处理 {results['successful_moves']} 个文件，删除源文件功能请在GUI界面使用")
            
            return results
        except Exception as e:
            raise FileOrganizerError(f"批量整理文件失败: {e}")
    def get_transfer_logs(self) -> List[str]:
        if not self.enable_transfer_log or not self.transfer_log_manager:
            raise FileOrganizerError("转移日志功能未启用")
        return self.transfer_log_manager.get_transfer_logs()
    def get_transfer_log_summary(self, log_file_path: str) -> Dict:
        if not self.enable_transfer_log or not self.transfer_log_manager:
            raise FileOrganizerError("转移日志功能未启用")
        return self.transfer_log_manager.get_session_summary(log_file_path)
    def restore_files_from_log(self, log_file_path: str, operation_ids: Optional[List[int]] = None, dry_run: bool = True) -> Dict:
        if not self.enable_transfer_log or not self.transfer_log_manager:
            raise FileOrganizerError("转移日志功能未启用")
        try:
            restore_results = self.transfer_log_manager.restore_from_log(
                log_file_path=log_file_path,
                operation_ids=operation_ids if operation_ids is not None else [],
                dry_run=dry_run
            )
            logging.info(f"文件恢复完成: 成功 {restore_results['successful_restores']}, 失败 {restore_results['failed_restores']}, 跳过 {restore_results['skipped_operations']}")
            return restore_results
        except Exception as e:
            raise FileOrganizerError(f"文件恢复失败: {e}")
    def cleanup_old_transfer_logs(self, days_to_keep: int = 30) -> int:
        if not self.enable_transfer_log or not self.transfer_log_manager:
            raise FileOrganizerError("转移日志功能未启用")
        return self.transfer_log_manager.cleanup_old_logs(days_to_keep)
    def _append_ai_result_to_file(self, ai_result_file: str, ai_result_item: dict) -> None:
        """追加AI结果到文件，避免在内存中保存大量数据"""
        try:
            # 读取现有数据
            existing_data = []
            if os.path.exists(ai_result_file):
                try:
                    with open(ai_result_file, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                        if content:  # 只有当文件不为空时才尝试解析JSON
                            existing_data = json.loads(content)
                        else:
                            existing_data = []  # 空文件时初始化为空列表
                except json.JSONDecodeError:
                    # 如果JSON解析失败，初始化为空列表
                    existing_data = []
                    logging.warning(f"AI结果文件格式错误，重新初始化: {ai_result_file}")
            
            # 追加新数据
            existing_data.append(ai_result_item)
            
            # 写回文件
            with open(ai_result_file, 'w', encoding='utf-8') as f:
                json.dump(existing_data, f, ensure_ascii=False, indent=2)
                
        except Exception as e:
            logging.error(f"追加AI结果到文件失败: {e}")
    
    def _update_last_ai_result(self, ai_result_file: str, updates: dict) -> None:
        """更新文件中最后一条AI结果记录"""
        try:
            # 读取现有数据
            existing_data = []
            if os.path.exists(ai_result_file):
                try:
                    with open(ai_result_file, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                        if content:  # 只有当文件不为空时才尝试解析JSON
                            existing_data = json.loads(content)
                        else:
                            existing_data = []  # 空文件时初始化为空列表
                except json.JSONDecodeError:
                    # 如果JSON解析失败，初始化为空列表
                    existing_data = []
                    logging.warning(f"AI结果文件格式错误，重新初始化: {ai_result_file}")
            
            # 更新最后一条记录
            if existing_data:
                existing_data[-1].update(updates)
                
                # 写回文件
                with open(ai_result_file, 'w', encoding='utf-8') as f:
                    json.dump(existing_data, f, ensure_ascii=False, indent=2)
                    
        except Exception as e:
            logging.error(f"更新AI结果文件失败: {e}")

    # def _ask_delete_source_files(self, successful_moves: List[Dict]) -> None:
    #     """询问用户是否删除源文件 - 已移至GUI界面，此方法不再使用"""
    #     # 删除源文件功能已移至GUI界面，不再在控制台询问
    #     pass

    def get_file_summary(self, file_path: str, max_length: int = 50, max_pages: int = 2, max_seconds: int = 10) -> str:
        """
        自动适配 txt/pdf/docx 格式，返回前max_length字摘要，PDF/Word只取前max_pages页，单文件处理超时max_seconds秒。
        """
        ext = Path(file_path).suffix.lower()
        start_time = time.time()
        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read(max_length)
            elif ext == '.pdf':
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    text = ''
                    for i, page in enumerate(reader.pages):
                        if i >= max_pages or (time.time() - start_time) > max_seconds:
                            break
                        page_text = page.extract_text() or ''
                        text += page_text
                        if len(text) >= max_length:
                            break
                    if (time.time() - start_time) > max_seconds:
                        return '提取超时，已跳过'
                    return text[:max_length] if text else '未能提取正文'
            elif ext == '.docx':
                doc = Document(file_path)
                text = ''
                for i, para in enumerate(doc.paragraphs):
                    if (time.time() - start_time) > max_seconds:
                        break
                    text += para.text
                    if len(text) >= max_length:
                        break
                if (time.time() - start_time) > max_seconds:
                    return '提取超时，已跳过'
                return text[:max_length] if text else '未能提取正文'
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read(max_length)
        except Exception as e:
            if (time.time() - start_time) > max_seconds:
                return '提取超时，已跳过'
            return f'摘要获取失败: {e}'
    
    def batch_read_documents(self, folder_path: str, progress_callback=None, summary_length: int = 200) -> Dict[str, Any]:
        """
        批量解读文件夹下所有文档，生成摘要并保存到AI结果文件
        
        Args:
            folder_path: 要解读的文件夹路径
            progress_callback: 进度回调函数，接收 (current, total, filename) 参数
            summary_length: 摘要长度（字数），默认200字
            
        Returns:
            包含处理结果的字典
        """
        start_time = time.time()
        timing_info = {}
        
        try:
            if not hasattr(self, 'ai_manager') or self.ai_manager is None:
                init_start = time.time()
                self.initialize_ollama()
                timing_info['ollama_init_time'] = round(time.time() - init_start, 3)
            
            folder_path = Path(folder_path)
            if not folder_path.exists() or not folder_path.is_dir():
                raise FileOrganizerError(f"文件夹不存在或不是有效目录: {folder_path}")
            
            # 支持的文档格式
            supported_extensions = {'.txt', '.pdf', '.docx', '.doc', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv'}
            
            # 扫描文件夹获取所有支持的文档文件
            document_files = []
            for file_path in folder_path.rglob('*'):
                if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                    document_files.append(file_path)
            
            if not document_files:
                return {
                    'total_files': 0,
                    'processed_files': 0,
                    'successful_reads': 0,
                    'failed_reads': 0,
                    'errors': [],
                    'timing_info': timing_info,
                    'start_time': datetime.now(),
                    'end_time': datetime.now()
                }
            
            # 初始化结果统计
            results = {
                'total_files': len(document_files),
                'processed_files': 0,
                'successful_reads': 0,
                'failed_reads': 0,
                'errors': [],
                'timing_info': timing_info,
                'start_time': datetime.now(),
                'end_time': None
            }
            
            # 初始化AI结果文件
            ai_result_file = 'ai_organize_result.json'
            if not os.path.exists(ai_result_file):
                with open(ai_result_file, 'w', encoding='utf-8') as f:
                    json.dump([], f, ensure_ascii=False, indent=2)
                logging.info(f"创建新的AI结果文件: {ai_result_file}")
            else:
                logging.info(f"AI结果文件已存在，将追加新记录: {ai_result_file}")
            

            
            logging.info(f"开始批量文档解读，共 {len(document_files)} 个文件")
            print(f"\n=== 开始批量文档解读 ===")
            print(f"源文件夹: {folder_path}")
            print(f"待处理文件总数: {len(document_files)}")
            print("=" * 50)
            
            # 逐个处理文档文件
            for i, file_path in enumerate(document_files, 1):
                filename = file_path.name
                
                # 调用进度回调
                if progress_callback:
                    progress_callback(i, len(document_files), filename)
                
                # 控制台输出当前处理进度
                print(f"\n[{i}/{len(document_files)}] 正在解读: {filename}")
                
                try:
                    logging.info(f"正在解读文件 {i}/{len(document_files)}: {filename}")
                    print(f"  📖 正在提取文件内容...", end="", flush=True)
                    
                    # 提取文件内容
                    extract_start = time.time()
                    extracted_content = self._extract_file_content(str(file_path))
                    extract_time = round(time.time() - extract_start, 3)
                    
                    # 根据设置截取内容用于AI处理
                    truncate_length = self.ai_parameters['content_extraction_length'] if self.ai_parameters['content_extraction_length'] < 2000 else len(extracted_content)
                    content_for_ai = extracted_content[:truncate_length] if extracted_content else ""
                    
                    print(f"\r  📖 正在生成摘要...", end="", flush=True)
                    
                    # 生成摘要
                    summary_start = time.time()
                    summary = self._generate_content_summary(content_for_ai, filename, summary_length)
                    summary_time = round(time.time() - summary_start, 3)
                    
                    print(f"\r  ✅ 解读完成")
                    print(f"     摘要: {summary[:100]}{'...' if len(summary) > 100 else ''}")
                    
                    # 构建AI结果项（与迁移功能格式兼容）
                    ai_result_item = {
                        "处理时间": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        "文件名": filename,
                        "源文件路径": str(file_path),  # 这里作为目标路径，因为文件已在目标位置
                        "文件摘要": summary,
                        "最匹配的目标目录": str(file_path.parent),  # 当前文件夹作为目标目录
                        "匹配理由": "批量文档解读",
                        "处理耗时信息": {
                            "总耗时(秒)": round(extract_time + summary_time, 3),
                            "内容提取耗时(秒)": extract_time,
                            "摘要生成耗时(秒)": summary_time,
                            "目录推荐耗时(秒)": 0
                        },
                        "最终目标路径": str(file_path),  # 文件当前位置就是最终位置
                        "操作类型": "文档解读",
                        "处理状态": "解读成功"
                    }
                    
                    # 保存到AI结果文件
                    self._append_ai_result_to_file(ai_result_file, ai_result_item)
                    
                    results['successful_reads'] += 1
                    logging.info(f"文档解读成功: {filename}，摘要长度: {len(summary)} 字符")
                    
                except Exception as e:
                    error_msg = f"解读文件 {filename} 时出错: {str(e)}"
                    print(f"\r  ❌ 解读失败: {str(e)}")
                    logging.error(error_msg)
                    results['errors'].append(error_msg)
                    results['failed_reads'] += 1
                    
                    # 保存失败记录
                    try:
                        ai_result_item = {
                            "处理时间": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                            "文件名": filename,
                            "源文件路径": str(file_path),
                            "文件摘要": "",
                            "最匹配的目标目录": str(file_path.parent),
                            "匹配理由": "批量文档解读",
                            "处理耗时信息": {
                                "总耗时(秒)": 0,
                                "内容提取耗时(秒)": 0,
                                "摘要生成耗时(秒)": 0,
                                "目录推荐耗时(秒)": 0
                            },
                            "最终目标路径": str(file_path),
                            "操作类型": "文档解读",
                            "处理状态": "解读失败",
                            "错误信息": str(e)
                        }
                        self._append_ai_result_to_file(ai_result_file, ai_result_item)
                    except Exception as save_error:
                        logging.warning(f"保存失败记录时出错: {save_error}")
                
                finally:
                    results['processed_files'] += 1
            
            results['end_time'] = datetime.now()
            
            # 输出处理完成总结
            print(f"\n=== 批量文档解读完成 ===")
            print(f"总文件数: {results['total_files']}")
            print(f"成功解读: {results['successful_reads']} 个")
            print(f"解读失败: {results['failed_reads']} 个")
            duration = (results['end_time'] - results['start_time']).total_seconds()
            print(f"总耗时: {duration:.1f} 秒")
            print("=" * 50)
            
            logging.info(f"批量文档解读完成: 成功 {results['successful_reads']}, 失败 {results['failed_reads']}")
            logging.info(f"解读结果已保存到: {ai_result_file}")
            
            return results
            
        except Exception as e:
            raise FileOrganizerError(f"批量文档解读失败: {e}")

    def _extract_folder_keywords(self, folder_path: str) -> List[str]:
        """
        从文件夹路径中提取有意义的关键词
        
        Args:
            folder_path: 文件夹路径
            
        Returns:
            List[str]: 关键词列表
        """
        try:
            # 提取路径中的文件夹名称
            path_parts = folder_path.split('\\') + folder_path.split('/')
            keywords = []
            
            for part in path_parts:
                part = part.strip()
                if not part:
                    continue
                
                # 移除编号前缀，如"【7-4-5】"
                if '【' in part and '】' in part:
                    part = part.split('】', 1)[1] if '】' in part else part
                
                # 过滤掉数字和单个字符
                if len(part) <= 1 or part.isdigit():
                    continue
                
                # 过滤掉包含连字符的数字代码，如"7-4-5"
                if '-' in part and all(segment.isdigit() for segment in part.split('-')):
                    continue
                
                # 过滤掉纯英文缩写
                if part.isupper() and len(part) <= 3:
                    continue
                
                # 避免重复添加
                if part not in keywords:
                    keywords.append(part)
                
                # 对于复合词汇，也提取其中的关键词
                if '\\' in part or '/' in part:
                    sub_parts = part.replace('\\', ' ').replace('/', ' ').split()
                    for sub_part in sub_parts:
                        sub_part = sub_part.strip()
                        if len(sub_part) > 1 and sub_part not in keywords:
                            keywords.append(sub_part)
            
            return keywords
        except Exception as e:
            logging.error(f"提取文件夹关键词失败: {e}")
            return []

    def _calculate_folder_relevance(self, file_content: str, file_summary: str, folder_path: str) -> float:
        """
        计算文件内容与文件夹路径的相关性评分
        
        Args:
            file_content: 文件内容
            file_summary: 文件摘要
            folder_path: 文件夹路径
            
        Returns:
            float: 相关性评分 (0.0-1.0)
        """
        try:
            # 提取文件夹路径中的关键词
            folder_keywords = self._extract_folder_keywords(folder_path)
            if not folder_keywords:
                return 0.0
            
            # 合并文件内容和摘要进行分析
            analysis_text = f"{file_content} {file_summary}".lower()
            
            # 定义通用的同义词和相关词汇映射（基于常见词汇）
            synonym_mapping = {
                # 通用业务词汇
                '综合': ['综合', '策略', '投资策略', '策略报告', '综合分析', '综合管理'],
                '研究': ['研究', '分析', '调研', '调查', '报告', '专题', '研究报告'],
                '投资': ['投资', '投资策略', '投资分析', '投资研究', '投资理念'],
                '行业': ['行业', '产业', '领域', '板块', '细分', '行业研究'],
                '公司': ['公司', '企业', '机构', '集团', '主体', '公司研究'],
                '经济': ['经济', '宏观经济', '经济分析', '经济研究', '经济模型'],
                '知识': ['知识', '百科', '科普', '教育', '学习'],
                '项目': ['项目', '合作', '合作项目', '项目合作'],
                '资料': ['资料', '文档', '文件', '材料', '信息'],
                '报告': ['报告', '研究报告', '分析报告', '策略报告', '专题报告'],
                '集合': ['集合', '汇总', '整理', '分类', '集合'],
                '分类': ['分类', '分类研究', '行业分类', '研究分类'],
                '权益': ['权益', '权益投资', '股票', '股权', '权益类'],
                '估值': ['估值', '估值方法', '估值研究', '价值评估'],
                '方法': ['方法', '研究方法', '分析方法', '估值方法'],
            }
            
            # 简化的相关性计算：主要基于关键词匹配和路径长度权重
            # 删除复杂的相斥性检查，让AI自己判断
            
            # 计算相关性评分
            total_score = 0.0
            matched_keywords = []
            
            for keyword in folder_keywords:
                keyword_lower = keyword.lower()
                
                # 直接匹配（完全匹配）
                if keyword_lower in analysis_text:
                    total_score += 1.0
                    matched_keywords.append(keyword)
                    continue
                
                # 部分匹配（关键词是文件内容中某个词汇的子串）
                if any(keyword_lower in word.lower() for word in analysis_text.split()):
                    total_score += 0.9
                    matched_keywords.append(f"{keyword}(部分匹配)")
                    continue
                
                # 同义词匹配
                if keyword in synonym_mapping:
                    synonyms = synonym_mapping[keyword]
                    for synonym in synonyms:
                        if synonym.lower() in analysis_text:
                            total_score += 0.8  # 同义词匹配权重稍低
                            matched_keywords.append(f"{keyword}(同义词:{synonym})")
                            break
            
            # 计算基础相关性评分
            if folder_keywords:
                base_relevance_score = total_score / len(folder_keywords)
                
                # 路径长度权重：路径越长（分类越细致），权重越高
                path_depth = len([p for p in folder_path.split('\\') + folder_path.split('/') if p.strip()])
                path_length_bonus = min(0.3, path_depth * 0.05)  # 最多增加0.3分
                
                # 最终相关性评分 = 基础评分 + 路径长度奖励
                relevance_score = base_relevance_score + path_length_bonus
                
                print(f"📊 相关性评分: {relevance_score:.3f} (基础: {base_relevance_score:.3f}, 路径奖励: {path_length_bonus:.3f}, 匹配关键词: {matched_keywords})")
                return relevance_score
            
            return 0.0
            
        except Exception as e:
            logging.error(f"计算文件夹相关性失败: {e}")
            return 0.0

    def _filter_irrelevant_folders(self, file_content: str, file_summary: str, recommended_paths: List[str], target_directory: str) -> List[str]:
        """
        过滤掉与文件内容不相关的文件夹路径，并按相关性评分排序
        
        Args:
            file_content: 文件内容
            file_summary: 文件摘要
            recommended_paths: AI推荐的文件夹路径列表
            target_directory: 目标目录
            
        Returns:
            List[str]: 按相关性评分排序的文件夹路径列表
        """
        try:
            print(f"🔍 开始相关性分析和排序...")
            print(f"📋 原始推荐路径: {recommended_paths}")
            
            relevance_scores = {}
            
            for path in recommended_paths:
                # 计算相关性评分
                relevance_score = self._calculate_folder_relevance(file_content, file_summary, path)
                relevance_scores[path] = relevance_score
                
                print(f"📊 {path}: 相关性评分 {relevance_score:.3f}")
            
            # 简化过滤逻辑：只过滤掉相关性评分为0的路径，保留所有有评分的路径
            filtered_paths = [path for path, score in relevance_scores.items() if score > 0]
            
            if not filtered_paths:
                print(f"⚠️ 所有路径相关性评分都为0，返回原始推荐")
                return recommended_paths  # 回退到原始推荐，而不是拒绝分类
            
            # 按相关性评分从高到低排序
            sorted_paths = sorted(filtered_paths, key=lambda x: relevance_scores[x], reverse=True)
            
            print(f"✅ 相关性排序完成:")
            for i, path in enumerate(sorted_paths, 1):
                print(f"  {i}. {path} (评分: {relevance_scores[path]:.3f})")
            
            return sorted_paths
            
        except Exception as e:
            logging.error(f"相关性过滤失败: {e}")
            print(f"❌ 相关性过滤失败: {e}")
            return recommended_paths

    def _extract_year_from_filename(self, filename: str) -> Optional[int]:
        """
        从文件名中提取年份信息
        支持多种年份格式：2024、2025、24、25等
        """
        import re
        
        # 提取4位年份 (2024, 2025等)
        year_patterns = [
            r'20\d{2}',  # 2020-2099
            r'19\d{2}',  # 1900-1999
        ]
        
        for pattern in year_patterns:
            matches = re.findall(pattern, filename)
            if matches:
                # 返回找到的第一个年份
                return int(matches[0])
        
        # 如果没有找到4位年份，尝试2位年份 (24, 25等)
        short_year_pattern = r'\b(2[0-9]|1[0-9])\b'  # 10-29
        matches = re.findall(short_year_pattern, filename)
        if matches:
            year = int(matches[0])
            # 假设20xx年
            if year >= 20:
                return 2000 + year
            else:
                return 1900 + year
        
        return None

    def _check_and_create_year_folder(self, target_folder_path: Path, filename: str) -> Path:
        """
        检查目标文件夹中是否包含文件年份，如果没有则创建年份子文件夹
        并参考原目录结构创建合适的子目录
        """
        # 提取文件名中的年份
        file_year = self._extract_year_from_filename(filename)
        
        if not file_year:
            print(f"📅 无法从文件名中提取年份: {filename}")
            return target_folder_path
        
        print(f"📅 从文件名提取到年份: {file_year}")
        
        # 检查目标文件夹中是否已有该年份的子文件夹
        year_folder_name = str(file_year)
        
        # 检查目标文件夹及其子文件夹中是否包含该年份
        existing_year_folders = []
        if target_folder_path.exists():
            for item in target_folder_path.iterdir():
                if item.is_dir():
                    # 检查文件夹名是否包含年份
                    if year_folder_name in item.name:
                        existing_year_folders.append(item)
                    # 递归检查子文件夹
                    for subitem in item.rglob("*"):
                        if subitem.is_dir() and year_folder_name in subitem.name:
                            existing_year_folders.append(subitem)
        
        if existing_year_folders:
            print(f"📅 找到现有年份文件夹: {[str(f) for f in existing_year_folders]}")
            # 使用第一个找到的年份文件夹
            return existing_year_folders[0]
        
        # 如果没有找到年份文件夹，需要创建新的年份文件夹结构
        # 分析原目录结构，创建合适的年份目录
        
        # 获取目标目录的根目录（通常是目标基础目录）
        # 例如：如果目标路径是 "F:\浩总参阅资料\【01】策略报告集合\2021\2021年投资策略\中信建投\行业"
        # 我们需要找到 "F:\浩总参阅资料\【01】策略报告集合" 作为根目录
        
        # 查找包含年份的目录层级
        target_parts = list(target_folder_path.parts)
        
        # 从后往前查找第一个包含年份的目录
        year_index = -1
        for i, part in enumerate(target_parts):
            if part.isdigit() and len(part) == 4 and part.startswith('20'):
                year_index = i
                break
        
        if year_index != -1:
            # 找到了年份目录，替换它
            new_parts = target_parts.copy()
            new_parts[year_index] = str(file_year)
            new_year_folder_path = Path(*new_parts)
        else:
            # 没有找到年份目录，在目标目录下创建年份文件夹
            new_year_folder_path = target_folder_path / year_folder_name
        
        print(f"📅 创建新的年份目录结构: {new_year_folder_path}")
        new_year_folder_path.mkdir(parents=True, exist_ok=True)
        
        return new_year_folder_path