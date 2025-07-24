#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件解读模块

本模块提供文件内容解读功能，包括：
1. 文件内容提取
2. 调用本地Ollama大模型进行内容分析
3. 生成文件内容摘要
4. 错误处理和日志记录

作者: AI Assistant
创建时间: 2025-01-15
"""

import os
import sys
import logging
import base64
import time
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List
from io import BytesIO

# 第三方库导入
try:
    import PyPDF2
except ImportError:
    print("请安装PyPDF2库: pip install PyPDF2")
    sys.exit(1)

try:
    import docx
except ImportError:
    print("请安装python-docx库: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
except ImportError:
    print("请安装Pillow库: pip install Pillow")
    sys.exit(1)

import json
from ai_client_manager import chat_with_ai


class FileReaderError(Exception):
    """文件解读异常类"""
    pass


class FileReader:
    """
    文件解读器类，使用统一的AI管理器
    """
    def __init__(self, model_name: Optional[str] = None, host: str = None):
        """
        初始化文件解读器
        Args:
            model_name: 指定使用的模型名称，如果为None则自动选择优先级模型
            host: 保留参数以兼容旧接口，实际使用统一的AI管理器
        """
        self.model_name = model_name
        self.summary_length = 200  # 默认摘要长度
        # 多模态模型列表（支持图像处理的模型）- 根据官方文档更新
        self.multimodal_models = ['llava', 'llava-llama3', 'llava-phi3', 'moondream', 
                                 'bakllava', 'llama3.2-vision', 'qwen2-vl', 'qwen-vl',
                                 'gemma3-vision', 'minicpm-v', 'cogvlm2', 'llava:7b', 
                                 'llava:13b', 'llava:34b', 'bakllava:7b', 'qwen3:0.6b',
                                 'deepseek-r1:8b']
        self.setup_logging()

    def setup_logging(self) -> None:
        """设置日志记录，仅输出到控制台"""
        # 配置日志格式
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.StreamHandler()
            ]
        )
        logging.info("文件解读器初始化完成")
    
    def extract_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取文件元数据（与智能文件分类器保持一致）"""
        try:
            file_path_obj = Path(file_path)
            stat = file_path_obj.stat()
            
            # 获取文件扩展名
            file_extension = file_path_obj.suffix.lower()
            
            # 获取文件大小（字节）
            file_size = stat.st_size
            
            # 获取创建时间和修改时间
            created_time = datetime.fromtimestamp(stat.st_ctime).isoformat()
            modified_time = datetime.fromtimestamp(stat.st_mtime).isoformat()
            
            return {
                'file_name': file_path_obj.name,
                'file_extension': file_extension,
                'file_size': file_size,
                'created_time': created_time,
                'modified_time': modified_time
            }
            
        except Exception as e:
            logging.error(f"提取文件元数据失败: {e}")
            return {
                'file_name': Path(file_path).name,
                'file_extension': Path(file_path).suffix.lower(),
                'file_size': 0,
                'created_time': '',
                'modified_time': ''
            }
    
    def extract_path_tags(self, file_path: str, base_folder: str = None) -> Dict[str, str]:
        """从文件路径中提取标签（去掉盘符和一级目录，从二级目录开始）"""
        import os
        from pathlib import Path
        tags = {}
        try:
            file_path_obj = Path(os.path.abspath(file_path)).resolve()
            file_parent = file_path_obj.parent
            
            # 获取完整的路径层级
            path_parts = list(file_parent.parts)
            
            # 去掉盘符和一级目录，从二级目录开始
            if len(path_parts) >= 2:
                # 从索引2开始（跳过盘符和一级目录）
                business_parts = path_parts[2:]
                
                for i, part in enumerate(business_parts):
                    tags[f"{i+1}级标签"] = part
                    
        except Exception as e:
            logging.error(f"提取路径标签失败: {e}")
        return tags
    


    def initialize_ollama(self) -> None:
        """
        兼容旧接口，使用统一的AI管理器
        """
        logging.info("AI客户端已通过统一管理器初始化")

    def extract_file_content(self, file_path: str, max_length: int = 2000) -> str:
        """
        提取文件内容
        
        Args:
            file_path: 文件路径
            max_length: 最大提取长度
            
        Returns:
            提取的文件内容
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileReaderError(f"文件不存在: {file_path}")
            
            file_extension = file_path.suffix.lower()
            logging.info(f"正在提取文件内容: {file_path.name} (类型: {file_extension})")
            
            # 根据文件类型选择不同的提取方法
            if file_extension == '.pdf':
                return self._extract_pdf_content(file_path, max_length)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_docx_content(file_path, max_length)
            elif file_extension in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv']:
                return self._extract_text_content(file_path, max_length)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
                return self._extract_image_info(file_path)
            else:
                # 尝试作为文本文件读取
                return self._extract_text_content(file_path, max_length)
                
        except Exception as e:
            logging.error(f"提取文件内容失败: {e}")
            return f"无法提取文件内容: {str(e)}"
    
    def _extract_pdf_content(self, file_path: Path, max_length: int) -> str:
        """提取PDF文件内容"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                
                # 读取前几页内容
                for page_num in range(min(3, len(pdf_reader.pages))):
                    page = pdf_reader.pages[page_num]
                    content += page.extract_text() + "\n"
                    
                    if len(content) >= max_length:
                        break
                
                return content[:max_length] if content else "PDF文件内容为空或无法提取"
                
        except Exception as e:
            return f"PDF文件读取失败: {str(e)}"
    
    def _extract_docx_content(self, file_path: Path, max_length: int) -> str:
        """提取Word文档内容"""
        try:
            # 检查文件是否存在且可读
            if not file_path.exists():
                return "文件不存在"
            
            if not file_path.is_file():
                return "路径不是文件"
            
            # 检查文件大小
            file_size = file_path.stat().st_size
            if file_size == 0:
                return "文件为空"
            
            # 尝试读取Word文档
            try:
                doc = docx.Document(file_path)
                content = ""
                
                # 提取段落内容
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # 跳过空段落
                        content += paragraph.text.strip() + "\n"
                        if len(content) >= max_length:
                            break
                
                # 如果段落内容为空，尝试提取表格内容
                if not content.strip():
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    content += cell.text.strip() + " "
                        content += "\n"
                        if len(content) >= max_length:
                            break
                
                return content[:max_length].strip() if content.strip() else "Word文档内容为空或无法提取"
                
            except Exception as docx_error:
                # 如果是.doc文件或损坏的.docx文件，尝试其他方法
                if file_path.suffix.lower() == '.doc':
                    return "不支持.doc格式，请转换为.docx格式"
                else:
                    return f"Word文档格式错误或文件损坏: {str(docx_error)}"
            
        except Exception as e:
            return f"Word文档读取失败: {str(e)}"
    
    def _extract_text_content(self, file_path: Path, max_length: int) -> str:
        """提取文本文件内容"""
        try:
            # 尝试多种编码格式
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        content = f.read(max_length)
                        
                        # 检查内容是否可读
                        printable_ratio = sum(1 for c in content if c.isprintable() or c.isspace()) / len(content) if content else 0
                        if printable_ratio > 0.7:
                            return content
                        
                except Exception:
                    continue
            
            return "文件内容为二进制格式，无法读取"
            
        except Exception as e:
            return f"文本文件读取失败: {str(e)}"
    
    def _extract_image_info(self, file_path: Path) -> str:
        """提取图片文件信息"""
        try:
            with Image.open(file_path) as img:
                info = f"图片文件信息:\n"
                info += f"格式: {img.format}\n"
                info += f"尺寸: {img.size[0]} x {img.size[1]}\n"
                info += f"模式: {img.mode}\n"
                
                # 如果有EXIF信息，提取一些基本信息
                if hasattr(img, '_getexif') and img._getexif():
                    info += "包含EXIF信息\n"
                
                return info
                
        except Exception as e:
            return f"图片文件信息提取失败: {str(e)}"
    
    def _image_to_base64(self, file_path: Path) -> str:
        """将图片转换为base64编码"""
        try:
            # 检查文件是否存在
            if not file_path.exists():
                logging.error(f"图片文件不存在: {file_path}")
                return ""
            
            # 检查文件大小（避免过大的图片）
            file_size = file_path.stat().st_size
            max_size = 10 * 1024 * 1024  # 10MB限制
            if file_size > max_size:
                logging.error(f"图片文件过大: {file_size} bytes > {max_size} bytes")
                return ""
            
            # 验证图片格式
            supported_formats = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']
            if file_path.suffix.lower() not in supported_formats:
                logging.error(f"不支持的图片格式: {file_path.suffix}")
                return ""
            
            # 尝试用PIL验证图片
            try:
                with Image.open(file_path) as img:
                    # 验证图片可以正常打开
                    img.verify()
            except Exception as img_error:
                logging.error(f"图片文件损坏或格式错误: {img_error}")
                return ""
            
            # 读取并编码图片
            with open(file_path, 'rb') as image_file:
                image_data = image_file.read()
                if not image_data:
                    logging.error("图片文件为空")
                    return ""
                
                encoded_string = base64.b64encode(image_data).decode('utf-8')
                logging.info(f"图片转base64成功，编码长度: {len(encoded_string)}")
                return encoded_string
                
        except Exception as e:
            logging.error(f"图片转base64失败: {e}")
            import traceback
            logging.error(f"详细错误信息: {traceback.format_exc()}")
            return ""
    
    def _is_multimodal_model(self, model_name: str) -> bool:
        """检测是否为多模态模型"""
        return any(mm_model in model_name.lower() for mm_model in self.multimodal_models)
    
    def _find_available_multimodal_model(self) -> str:
        """查找可用的多模态模型"""
        # 使用统一的AI管理器获取可用模型
        try:
            from ai_client_manager import get_available_models
            available_models = get_available_models()
            for model in available_models:
                if self._is_multimodal_model(model):
                    return model
        except Exception as e:
            logging.error(f"获取可用模型失败: {e}")
        return None
    
    def _process_image_with_multimodal(self, file_path: str, max_summary_length: int) -> str:
        """使用多模态模型处理图像文件"""
        try:
            # 检查文件是否存在
            file_path_obj = Path(file_path)
            if not file_path_obj.exists():
                return "图像文件不存在"
            
            # 将图像转换为base64编码
            image_base64 = self._image_to_base64(file_path_obj)
            if not image_base64:
                return "图像文件读取失败或转换base64失败"
            
            # 构建多模态提示词
            prompt = f"请仔细分析这张图片的内容，并生成一个不超过{max_summary_length}字的详细摘要。请描述图片中的主要内容、对象、场景、文字信息等关键信息。如果是文档图片，请提取其中的文字内容。/no_think"
            
            # 调用多模态模型 - 按照官方文档格式
            logging.info(f"使用多模态模型 {self.model_name} 分析图像: {file_path_obj.name}")
            
            # 构建消息，确保images参数正确传递
            messages = [
                {
                    'role': 'user',
                    'content': prompt,
                    'images': [image_base64]  # base64编码的图像数据
                }
            ]
            
            # 使用统一的AI管理器
            response = self._chat_with_retry(messages)
            
            # 检查响应格式
            if response and isinstance(response, str):
                content = response.strip()
                if content:
                    # 立即清理AI响应中的<think>标签
                    content = self._clean_ai_response(content)
                    logging.info(f"多模态模型分析成功，内容长度: {len(content)}")
                    return content
                else:
                    return "多模态模型返回空内容"
            else:
                logging.error(f"多模态模型响应类型异常: {type(response)}")
                return "多模态模型响应类型异常"
                
        except Exception as e:
            logging.error(f"多模态模型处理图像失败: {e}")
            import traceback
            logging.error(f"详细错误信息: {traceback.format_exc()}")
            return f"图像分析失败: {str(e)}"
    
    def generate_summary(self, file_path: str, max_summary_length: int = 50) -> Dict[str, Any]:
        """
        生成文件内容摘要（与智能文件分类器保持一致）
        
        Args:
            file_path: 文件路径
            max_summary_length: 摘要最大长度
            
        Returns:
            包含摘要信息的字典，包括原始文本和摘要
        """
        start_time = time.time()
        timing_info = {}
        
        result = {
            'file_path': file_path,
            'file_name': Path(file_path).name,
            'success': False,
            'extracted_text': '',  # 新增：提取的原始文本
            'summary': '',
            'error': '',
            'model_used': self.model_name,
            'timestamp': datetime.now().isoformat(),
            'timing_info': timing_info
        }
        
        try:
            # 第一步：提取文件元数据
            metadata_start = time.time()
            file_metadata = self.extract_file_metadata(file_path)
            metadata_time = round(time.time() - metadata_start, 3)
            timing_info['metadata_extraction_time'] = metadata_time
            logging.info(f"文件元数据提取完成，耗时: {metadata_time}秒")
            
            # 初始化AI客户端（如果尚未初始化）
            if not hasattr(self, 'client_initialized'):
                self.initialize_ollama()
                self.client_initialized = True
            
            file_path_obj = Path(file_path)
            file_ext = file_path_obj.suffix.lower()
            
            logging.info(f"开始解读文件: {file_path_obj.name}")
            
            # 检查是否为图像文件
            image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']
            is_image = file_ext in image_extensions
            
            if is_image:
                # 图像文件处理
                if self._is_multimodal_model(self.model_name):
                    # 使用多模态模型处理图像
                    logging.info(f"使用多模态模型 {self.model_name} 处理图像文件")
                    summary = self._process_image_with_multimodal(file_path, max_summary_length)
                    result['extracted_text'] = "[图像文件，无文本内容]"
                else:
                    # 如果当前模型不支持多模态，尝试切换到多模态模型
                    multimodal_model = self._find_available_multimodal_model()
                    if multimodal_model:
                        logging.info(f"切换到多模态模型 {multimodal_model} 处理图像文件")
                        old_model = self.model_name
                        self.model_name = multimodal_model
                        summary = self._process_image_with_multimodal(file_path, max_summary_length)
                        result['extracted_text'] = "[图像文件，无文本内容]"
                    else:
                        # 没有可用的多模态模型，提取图像基本信息
                        logging.warning("没有可用的多模态模型，仅提取图像基本信息")
                        image_info = self._extract_image_info(file_path_obj)
                        summary = f"图像文件基本信息: {image_info}"
                        result['extracted_text'] = "[图像文件，无文本内容]"
            else:
                # 非图像文件，先提取文本内容
                logging.info("正在提取文件文本内容...")
                file_content = self.extract_file_content(file_path)
                result['extracted_text'] = file_content  # 保存提取的原始文本
                
                # 输出提取的文本内容到日志
                logging.info(f"提取的文本内容（前200字符）: {file_content[:200]}...")
                print(f"\n=== 提取的原始文本内容 ===")
                print(f"文件: {file_path_obj.name}")
                print(f"文本长度: {len(file_content)} 字符")
                print(f"内容预览:\n{file_content[:500]}{'...' if len(file_content) > 500 else ''}")
                print("=" * 50)
                
                # 构建摘要提示词并调用模型
                prompt = self._build_summary_prompt(file_content, max_summary_length)
                
                logging.info("正在调用大模型生成摘要...")
                print(f"\n=== 正在生成摘要 ===")
                print(f"使用模型: {self.model_name}")
                print(f"摘要长度限制: {max_summary_length} 字符")
                
                # 使用系统提示词来抑制思考过程
                messages = [
                    {
                        'role': 'system',
                        'content': '你是一个专业的文档摘要助手。重要：不要输出任何推理过程、思考步骤或解释。直接按要求输出结果。只输出摘要内容，不要包含任何其他信息。'
                    },
                    {
                        'role': 'user',
                        'content': prompt
                    }
                ]
                
                summary = self._chat_with_retry(messages)
                
                # 输出生成的摘要
                print(f"\n=== 生成的摘要 ===")
                print(f"摘要内容: {summary}")
                print(f"摘要长度: {len(summary)} 字符")
                print("=" * 50)
            
            # 处理摘要结果
            summary = summary.strip()
            
            # 清理可能的思考过程标签和内容
            summary = summary.replace('<think>', '').replace('</think>', '').strip()
            
            # 移除常见的思考过程开头
            think_prefixes = [
                '好的，', '好，', '嗯，', '我来', '我需要', '首先，', '让我', '现在我要',
                '用户希望', '用户要求', '用户让我', '根据', '基于', '考虑到', '让我先仔细看看',
                '用户给了我这个查询', '用户给了我这个任务', '用户给了一个任务',
                '首先，我得看一下', '首先，我要理解', '首先，我得仔细看看',
                '好的，用户让我', '用户让我生成', '内容来自文件', '重点包括', '首先，我需要确认'
            ]
            
            # 更激进的清理：移除所有以思考过程开头的句子
            lines = summary.split('\n')
            cleaned_lines = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 检查是否以思考过程开头
                is_think_line = False
                for prefix in think_prefixes:
                    if line.lower().startswith(prefix.lower()):
                        is_think_line = True
                        break
                
                # 如果不是思考过程，保留这一行
                if not is_think_line:
                    cleaned_lines.append(line)
            
            # 如果清理后没有内容，尝试更简单的方法
            if not cleaned_lines:
                # 找到第一个不是思考过程的句子
                sentences = summary.split('。')
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                    
                    # 检查是否以思考过程开头
                    is_think_sentence = False
                    for prefix in think_prefixes:
                        if sentence.lower().startswith(prefix.lower()):
                            is_think_sentence = True
                            break
                    
                    if not is_think_sentence:
                        cleaned_lines.append(sentence)
                        break
            
            # 如果还是没有内容，使用原始摘要的最后一部分
            if not cleaned_lines:
                # 取最后100个字符作为摘要
                summary = summary[-100:] if len(summary) > 100 else summary
            
            # 重新组合清理后的内容
            if cleaned_lines:
                summary = '。'.join(cleaned_lines)
            
            # 确保摘要长度不超过限制
            if len(summary) > max_summary_length:
                summary = summary[:max_summary_length-3] + "..."
            
            # 计算总处理时间
            total_time = round(time.time() - start_time, 3)
            timing_info['total_processing_time'] = total_time
            
            result['success'] = True
            result['summary'] = summary
            result['model_used'] = self.model_name
            result['file_metadata'] = file_metadata
            result['timing_info'] = timing_info
            
            logging.info(f"文件解读成功: {file_path_obj.name} -> {summary[:50]}...，总耗时: {total_time}秒")
            
        except Exception as e:
            error_msg = f"文件解读失败: {str(e)}"
            result['error'] = error_msg
            logging.error(error_msg)
        
        return result
    
    def _build_summary_prompt(self, file_content: str, max_length: int) -> str:
        """构建摘要生成提示词"""
        prompt = f"""请为以下文件内容生成一个{max_length}字以内的中文摘要。

文件内容：
{file_content}

要求：
1. 概括文件的主要内容和主题
2. 突出关键信息和要点
3. 语言简洁明了
4. 字数控制在{max_length}字以内
5. 直接输出摘要内容，不要包含任何思考过程或说明文字
6. 不要使用"<think>"标签或任何思考过程描述

摘要：/no_think"""
        return prompt
    
    def _clean_ai_response(self, response: str) -> str:
        """清理AI响应中的思考过程"""
        if not response:
            return response
        
        # 清理可能的思考过程标签和内容
        response = response.replace('<think>', '').replace('</think>', '').strip()
        
        # 移除常见的思考过程开头
        think_prefixes = [
            '好的，', '好，', '嗯，', '我来', '我需要', '首先，', '让我', '现在我要',
            '用户希望', '用户要求', '用户让我', '根据', '基于', '考虑到', '让我先仔细看看',
            '用户给了我这个查询', '用户给了我这个任务', '用户给了一个任务',
            '首先，我得看一下', '首先，我要理解', '首先，我得仔细看看',
            '好的，用户让我', '用户让我生成', '内容来自文件', '重点包括', '首先，我需要确认'
        ]
        
        # 更激进的清理：移除所有以思考过程开头的句子
        lines = response.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检查是否以思考过程开头
            is_think_line = False
            for prefix in think_prefixes:
                if line.lower().startswith(prefix.lower()):
                    is_think_line = True
                    break
            
            # 如果不是思考过程，保留这一行
            if not is_think_line:
                cleaned_lines.append(line)
        
        # 如果清理后没有内容，尝试更简单的方法
        if not cleaned_lines:
            # 找到第一个不是思考过程的句子
            sentences = response.split('。')
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                
                # 检查是否以思考过程开头
                is_think_sentence = False
                for prefix in think_prefixes:
                    if sentence.lower().startswith(prefix.lower()):
                        is_think_sentence = True
                        break
                
                if not is_think_sentence:
                    cleaned_lines.append(sentence)
                    break
        
        # 如果还是没有内容，使用原始响应的最后一部分
        if not cleaned_lines:
            # 取最后100个字符
            response = response[-100:] if len(response) > 100 else response
            return response
        
        # 重新组合清理后的内容
        if cleaned_lines:
            response = '\n'.join(cleaned_lines)
        
        return response.strip()
    
    def _chat_with_retry(self, messages: list, max_retries: int = 3, images: List[str] = None) -> str:
        """
        使用统一的AI管理器进行AI调用
        Args:
            messages: 消息列表
            max_retries: 最大重试次数（保留参数以兼容旧接口）
            images: base64编码的图像列表（用于多模态模型）
        Returns:
            模型响应内容
        Raises:
            FileReaderError: 当所有重试都失败时抛出
        """
        try:
            # 如果有图像数据，添加到消息中
            if images and messages and len(messages) > 0:
                messages[0]['images'] = images
            
            # 使用统一的AI管理器
            response = chat_with_ai(messages)
            
            # 清理AI响应中的<think>标签
            content = self._clean_ai_response(response)
            return content
            
        except Exception as e:
            raise FileReaderError(f"AI调用失败: {e}")
    
    def get_available_models(self) -> list:
        """
        获取可用模型列表
        
        Returns:
            可用模型名称列表
        """
        try:
            from ai_client_manager import get_available_models
            return get_available_models()
        except Exception as e:
            logging.error(f"获取模型列表失败: {e}")
            return []
    
    def append_result_to_file(self, ai_result_file: str, result: dict, base_folder: str = "") -> None:
        """将文件解读结果追加到JSON文件（与智能文件分类器保持一致）"""
        try:
            # 构建结果条目
            entry = {
                "处理时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "文件名": result['file_name'],
                "文件摘要": result['summary'],
                "最匹配的目标目录": "",  # 文件解读功能不涉及目录匹配
                "处理耗时": result['timing_info'].get('total_processing_time', 0),
                "最终目标路径": result['file_path'],  # 文件当前存储的完整路径
                "操作类型": "文件解读",
                "处理状态": "解读成功" if result['success'] else "解读失败",
                "标签": result.get('tags', {}),
                "文件元数据": {
                    "file_name": result['file_metadata']['file_name'],
                    "file_extension": result['file_metadata']['file_extension'],
                    "file_size": result['file_metadata']['file_size'],
                    "created_time": result['file_metadata']['created_time'],
                    "modified_time": result['file_metadata']['modified_time']
                }
            }
            
            # 读取现有文件
            existing_data = []
            if os.path.exists(ai_result_file):
                try:
                    with open(ai_result_file, 'r', encoding='utf-8') as f:
                        existing_data = json.load(f)
                except (json.JSONDecodeError, FileNotFoundError):
                    existing_data = []
            
            # 添加新条目
            existing_data.append(entry)
            
            # 写入文件
            with open(ai_result_file, 'w', encoding='utf-8') as f:
                json.dump(existing_data, f, ensure_ascii=False, indent=2)
            
            logging.info(f"文件解读结果已写入: {ai_result_file}")
            
        except Exception as e:
            logging.error(f"写入结果文件失败: {e}")
